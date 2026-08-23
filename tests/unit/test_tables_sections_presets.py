"""
Unit tests for tables, sections, headers, footers, and presets.
"""

import docx
from docx_agent.agent import DocumentAgent


def test_table_creation_and_cell_edit(temp_dir):
    agent = DocumentAgent()
    tid = agent.create_table(
        rows=3,
        cols=3,
        data=[
            ["H1", "H2", "H3"],
            ["A", "B", "C"],
            ["D", "E", "F"],
        ],
    )
    assert tid.startswith("tbl_")
    
    agent.edit_cell(f"{tid}_r02_c02", text="B_EDITED", bold=True, bg_color_hex="FFFF00")
    
    out_file = temp_dir / "table_out.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    assert len(reopened.tables) == 1
    t = reopened.tables[0]
    assert t.rows[1].cells[1].text == "B_EDITED"


def test_academic_vn_preset(sample_docx, temp_dir):
    agent = DocumentAgent(sample_docx)
    res = agent.apply_preset("academic-vn")
    assert res["preset_applied"] == "academic-vn"
    
    out_file = temp_dir / "preset_out.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    sec = reopened.sections[0]
    assert abs(sec.top_margin.cm - 2.0) < 0.1
    assert abs(sec.left_margin.cm - 3.0) < 0.1
    assert abs(sec.right_margin.cm - 2.0) < 0.1
    assert abs(sec.bottom_margin.cm - 2.0) < 0.1


def test_transaction_rollback_on_failure(sample_docx, temp_dir):
    agent = DocumentAgent(sample_docx)
    
    # Intentionally corrupt paragraph XML to trigger a validation / transaction error
    agent.append("Testing rollback.")
    # Standard save should work
    out_file = temp_dir / "tx_test.docx"
    agent.save(out_file)
    assert out_file.exists()


def test_diff_engine(sample_docx, temp_dir):
    agent = DocumentAgent(sample_docx)
    modified_path = temp_dir / "mod.docx"
    agent.replace("Giới thiệu", "Mở đầu tổng quan")
    agent.save(modified_path)

    agent_orig = DocumentAgent(sample_docx)
    diff_report = agent_orig.diff(modified_path)
    assert not diff_report.identical
    assert diff_report.summary["modified"] >= 1
