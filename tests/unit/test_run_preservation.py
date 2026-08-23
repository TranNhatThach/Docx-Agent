"""
Format Preservation Tests: Verifies that run surgeries preserve styling (bold, italic, colors).
"""

import docx
from docx.shared import RGBColor
from docx_agent.agent import DocumentAgent


def test_replace_inside_single_run_preserves_formatting(mixed_runs_docx, temp_dir):
    agent = DocumentAgent(mixed_runs_docx)
    
    # Target 'world' inside [Hello ] [BOLD+RED(world)] [ITALIC( today and welcome!)]
    n = agent.replace(target="world", replacement="universe")
    assert n >= 1
    
    out_file = temp_dir / "preserved_out.docx"
    agent.save(out_file)

    # Reopen independently and verify
    reopened = docx.Document(str(out_file))
    p = reopened.paragraphs[0]
    
    assert "Hello universe today and welcome!" in p.text
    
    # Verify that 'universe' has bold = True
    found_universe_bold = False
    for r in p.runs:
        if "universe" in r.text:
            assert r.bold is True
            found_universe_bold = True
            
    assert found_universe_bold, "Run containing 'universe' must retain bold formatting!"


def test_replace_spanning_across_runs(mixed_runs_docx, temp_dir):
    agent = DocumentAgent(mixed_runs_docx)
    
    # Paragraph 2 has runs: ['hel', 'lo universe'] -> target 'hello universe' spans both runs!
    n = agent.replace(target="hello universe", replacement="welcome everyone")
    assert n >= 1
    
    out_file = temp_dir / "cross_runs_out.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    p2 = reopened.paragraphs[1]
    
    assert "welcome everyone" in p2.text
    assert "hel" not in p2.text
