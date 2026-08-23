"""
Unit tests for MCP tool dispatching and Markdown to DOCX converter.
"""

from pathlib import Path
import docx
from docx_agent.interfaces.mcp.server import handle_tool_call, MCP_TOOLS
from docx_agent.operations.markdown import MarkdownToDocxConverter


def test_mcp_tools_list_schema():
    assert len(MCP_TOOLS) >= 12
    tool_names = [t["name"] for t in MCP_TOOLS]
    assert "docx_inspect" in tool_names
    assert "docx_replace" in tool_names
    assert "docx_preset" in tool_names
    assert "docx_verify" in tool_names
    assert "docx_diff" in tool_names


def test_mcp_inspect_dispatch(sample_docx):
    res = handle_tool_call("docx_inspect", {"file": str(sample_docx)})
    assert "paragraphs_count" in res
    assert res["paragraphs_count"] >= 4


def test_mcp_replace_dispatch(sample_docx, temp_dir):
    out_file = temp_dir / "mcp_replace_out.docx"
    res = handle_tool_call("docx_replace", {
        "file": str(sample_docx),
        "target": "Giới thiệu",
        "replacement": "Khái Quát",
        "output": str(out_file),
    })
    assert res["success"] is True
    assert res["replaced_count"] >= 1
    assert Path(out_file).exists()


def test_markdown_to_docx_conversion(temp_dir):
    md_content = """# Tiêu đề Luận Văn

## 1. Mở Đầu
Đây là văn bản mở đầu với các điểm nhấn.

- Điểm thứ nhất
- Điểm thứ hai

| Cột 1 | Cột 2 |
|---|---|
| Dữ liệu A | Dữ liệu B |

```python
def hello_world():
    print("Hello from code")
```
"""
    md_file = temp_dir / "doc.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(md_content)

    docx_out = temp_dir / "converted.docx"
    saved = MarkdownToDocxConverter.convert(md_file, docx_out, preset="academic-vn")
    
    assert Path(saved).exists()
    doc = docx.Document(saved)
    assert len(doc.paragraphs) >= 4
    assert len(doc.tables) == 1
    assert "Tiêu đề Luận Văn" in doc.paragraphs[0].text
