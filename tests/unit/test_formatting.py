"""
Unit tests for text formatting and paragraph styling operations.
"""

import docx
from docx.shared import Pt, Cm
from docx_agent.agent import DocumentAgent


def test_format_text_properties(sample_docx, temp_dir):
    agent = DocumentAgent(sample_docx)
    
    agent.format_text(
        target="p_0002",
        font_name="Arial",
        font_size_pt=14.0,
        bold=True,
        italic=True,
        color_rgb="003366",
    )
    
    out_file = temp_dir / "formatted_text.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    p = reopened.paragraphs[1]
    assert len(p.runs) > 0
    assert p.runs[0].font.name == "Arial"
    assert p.runs[0].font.size.pt == 14.0
    assert p.runs[0].bold is True
    assert p.runs[0].italic is True


def test_format_paragraph_properties(sample_docx, temp_dir):
    agent = DocumentAgent(sample_docx)
    
    agent.format_paragraph(
        target="p_0003",
        alignment="justify",
        line_spacing=1.5,
        space_before_pt=10.0,
        space_after_pt=12.0,
        first_line_indent_cm=1.27,
    )
    
    out_file = temp_dir / "formatted_para.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    p = reopened.paragraphs[2]
    fmt = p.paragraph_format
    assert abs(fmt.line_spacing - 1.5) < 0.01
    assert abs(fmt.first_line_indent.cm - 1.27) < 0.05
