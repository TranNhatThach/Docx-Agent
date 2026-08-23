"""
Fixture Generator for Docx-Agent Test Suite.
Generates realistic .docx files: simple, Vietnamese, academic_report, tables,
images, multi_page, complex_formatting, unsupported_ooxml.
"""

from pathlib import Path
import docx
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FIXTURES_DIR = Path(__file__).parent


def create_simple_docx() -> Path:
    p = FIXTURES_DIR / "simple.docx"
    doc = docx.Document()
    doc.add_heading("Tài Liệu Cơ Bản", level=1)
    doc.add_paragraph("Đây là đoạn văn đầu tiên trong tài liệu đơn giản.")
    doc.add_paragraph("Đoạn văn thứ hai chứa thông tin tổng quát.")
    doc.save(str(p))
    return p


def create_vietnamese_docx() -> Path:
    p = FIXTURES_DIR / "Vietnamese.docx"
    doc = docx.Document()
    
    # Page setup
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)

    # Title
    h0 = doc.add_heading("BÁO CÁO NGHIÊN CỨU KHOA HỌC", level=0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1
    doc.add_heading("1. Đặt vấn đề và tính cấp thiết", level=1)
    doc.add_paragraph(
        "Trong bối cảnh chuyển đổi số và phát triển trí tuệ nhân tạo hiện nay, việc quản lý và biên tập "
        "các tài liệu học thuật theo định dạng Microsoft Word (.docx) đặt ra nhiều thách thức về tính toàn vẹn, "
        "chuẩn mực định dạng theo Tiêu chuẩn Quốc gia Việt Nam (TCVN) và khả năng cộng tác thông minh."
    )

    # Heading 2
    doc.add_heading("1.1 Mục tiêu nghiên cứu", level=2)
    doc.add_paragraph(
        "Mục tiêu chính của đề tài là xây dựng hệ thống biên tập tài liệu kết hợp tác tử AI thế hệ mới, "
        "đảm bảo giữ nguyên 100% định dạng nguyên bản, hỗ trợ kiểm tra quy chuẩn và tích hợp trích dẫn khoa học chính xác."
    )

    # Heading 2
    doc.add_heading("1.2 Đối tượng và phạm vi", level=2)
    doc.add_paragraph(
        "Đối tượng nghiên cứu bao gồm các tài liệu báo cáo kỹ thuật, luận văn thạc sĩ, đồ án tốt nghiệp "
        "được soạn thảo bằng tiếng Việt với bảng mã Unicode UTF-8 chuẩn."
    )

    doc.save(str(p))
    return p


def create_tables_docx() -> Path:
    p = FIXTURES_DIR / "tables.docx"
    doc = docx.Document()
    doc.add_heading("Báo Cáo Thử Nghiệm Bảng Biểu", level=1)
    doc.add_paragraph("Bảng 1 dưới đây tổng hợp kết quả đánh giá hiệu năng của hệ thống:")

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    headers = ["STT", "Thành phần", "Thời gian phản hồi (ms)", "Độ chính xác (%)"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    data = [
        ["1", "Bộ phân tích OpenXML", "45.2", "99.8%"],
        ["2", "Mô hình tài liệu Canonical", "12.6", "100.0%"],
        ["3", "Công cụ giao dịch Agent", "28.4", "99.5%"],
    ]
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx + 1].cells[c_idx].text = val

    doc.save(str(p))
    return p


def create_complex_formatting_docx() -> Path:
    p = FIXTURES_DIR / "complex_formatting.docx"
    doc = docx.Document()
    doc.add_heading("Định Dạng Phức Tạp và Đa Dạng", level=1)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("Văn bản bình thường, ")
    r2 = p1.add_run("văn bản in đậm, ")
    r2.bold = True
    r3 = p1.add_run("văn bản in nghiêng, ")
    r3.italic = True
    r4 = p1.add_run("văn bản gạch chân, ")
    r4.underline = True
    r5 = p1.add_run("văn bản màu đỏ (#FF0000).")
    r5.font.color.rgb = RGBColor(255, 0, 0)

    p2 = doc.add_paragraph()
    p2.add_run("Công thức hóa học: H")
    r_sub = p2.add_run("2")
    r_sub.font.subscript = True
    p2.add_run("O và phương trình năng lượng: E = mc")
    r_sup = p2.add_run("2")
    r_sup.font.superscript = True

    doc.save(str(p))
    return p


def create_multi_page_docx() -> Path:
    p = FIXTURES_DIR / "multi_page.docx"
    doc = docx.Document()
    
    for ch in range(1, 11):
        doc.add_heading(f"Chương {ch}: Phân Tích Chuyên Sâu Phần {ch}", level=1)
        for sub in range(1, 4):
            doc.add_heading(f"{ch}.{sub} Nội dung chi tiết phần {ch}.{sub}", level=2)
            for _ in range(3):
                doc.add_paragraph(
                    f"Đây là đoạn văn thứ nhất trong mục {ch}.{sub}. Hệ thống phân tích tài liệu tự động "
                    f"cần đảm bảo việc phân trang chính xác trên từng trang A4, hỗ trợ căn lề chuẩn, giãn dòng 1.5, "
                    f"và đánh số trang liên tục từ đầu đến cuối văn bản mà không bị gián đoạn."
                )
        if ch < 10:
            doc.add_page_break()

    doc.save(str(p))
    return p


def create_academic_report_docx() -> Path:
    p = FIXTURES_DIR / "academic_report.docx"
    doc = docx.Document()

    # Section properties
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)

    # Title
    t = doc.add_heading("BÁO CÁO KIẾN TRÚC HỆ THỐNG DOCX-AGENT V2.1", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Tổng Quan Kiến Trúc", level=1)
    doc.add_paragraph(
        "Hệ thống Docx-Agent V2.1 được thiết kế theo mô hình phân lớp gồm ba tầng chính: "
        "Tầng Trình soạn thảo A4 (Document Workspace), Tầng Mô hình dữ liệu chuẩn tắc (Canonical Document Model), "
        "và Tầng Cầu nối Tác tử AI (Antigravity Agent Bridge)."
    )

    doc.add_heading("2. Các Thành Phần Trọng Tâm", level=1)
    doc.add_heading("2.1 Bộ Xử Lý Giao Dịch Agent (Transaction Manager)", level=2)
    doc.add_paragraph(
        "Mọi đề xuất chỉnh sửa từ Antigravity Agent đều được đóng gói thành các giao dịch nguyên tử (Atomic Transactions). "
        "Người dùng có toàn quyền xem trước (Preview Diff), chấp thuận (Apply) hoặc từ chối (Reject) trước khi thay đổi được ghi vào tài liệu."
    )

    doc.add_heading("2.2 Trích Dẫn Học Thuật và Nguồn Dữ Liệu", level=2)
    doc.add_paragraph(
        "Hệ thống hỗ trợ kiểm chứng nguồn tài liệu theo tiêu chuẩn APA, IEEE và TCVN, đảm bảo thông tin trích dẫn luôn xác thực "
        "và không tạo ra các tài liệu tham khảo ảo."
    )

    # Table
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Mô-đun"
    tbl.rows[0].cells[1].text = "Chức năng"
    tbl.rows[0].cells[2].text = "Độ tin cậy"

    tbl.rows[1].cells[0].text = "Canonical Engine"
    tbl.rows[1].cells[1].text = "Lưu trữ in-memory chuẩn hóa"
    tbl.rows[1].cells[2].text = "100%"

    tbl.rows[2].cells[0].text = "Visual Verifier"
    tbl.rows[2].cells[1].text = "Kiểm tra bố cục và tràn trang"
    tbl.rows[2].cells[2].text = "99.9%"

    doc.save(str(p))
    return p


def create_unsupported_ooxml_docx() -> Path:
    p = FIXTURES_DIR / "unsupported_ooxml.docx"
    doc = docx.Document()
    doc.add_heading("Kiểm Tra Đối Tượng OOXML Đặc Biệt", level=1)
    doc.add_paragraph("Đoạn văn tiêu chuẩn trước phần tử đặc biệt.")

    # Inject a custom bookmark or sdt XML node to simulate advanced Word features
    custom_xml = '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:sdtPr><w:id w:val="998877"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>Nội dung trong Structured Document Tag</w:t></w:r></w:p></w:sdtContent></w:sdt>'
    sdt_elem = parse_xml(custom_xml)
    doc._body._element.append(sdt_elem)

    doc.add_paragraph("Đoạn văn sau phần tử đặc biệt.")
    doc.save(str(p))
    return p


def create_images_docx() -> Path:
    import struct
    import zlib

    # Generate a tiny valid PNG file
    def generate_minimal_png(width=100, height=100) -> bytes:
        def chunk(tag, data):
            return struct.pack("!I", len(data)) + tag + data + struct.pack("!I", zlib.crc32(tag + data) & 0xffffffff)
        header = b"\x89PNG\r\n\x1a\n"
        ihdr = chunk(b"IHDR", struct.pack("!IIBBBBB", width, height, 8, 2, 0, 0, 0))
        raw_data = b"".join(b"\x00" + b"\x25\x63\xeb" * width for _ in range(height))
        idat = chunk(b"IDAT", zlib.compress(raw_data))
        iend = chunk(b"IEND", b"")
        return header + ihdr + idat + iend

    img_path = FIXTURES_DIR / "sample_test_image.png"
    with open(img_path, "wb") as f:
        f.write(generate_minimal_png(200, 100))

    p = FIXTURES_DIR / "images.docx"
    doc = docx.Document()
    doc.add_heading("Tài Liệu Chứa Hình Ảnh và Chú Thích", level=1)
    doc.add_paragraph("Dưới đây là hình ảnh kiến trúc mẫu minh họa cho hệ thống:")
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(img_path), width=Cm(12.0))
    
    p_cap = doc.add_paragraph("Hình 1: Kiến trúc tổng thể Docx-Agent V2.1", style="Caption")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(p))
    return p


def create_corrupted_docx() -> Path:
    p = FIXTURES_DIR / "corrupted.docx"
    with open(p, "wb") as f:
        # PK header with completely truncated corrupted zip stream
        f.write(b"PK\x03\x04\x14\x00\x00\x00\x08\x00\x00\x00CORRUPTED_DOCX_PAYLOAD_TEST_FAILURES_1234567890")
    return p


def create_stress_50_page_docx() -> Path:
    p = FIXTURES_DIR / "stress_50_page.docx"
    doc = docx.Document()
    
    # 50 chapters, each having 3 subheadings and multiple paragraphs
    for chap_idx in range(1, 51):
        doc.add_heading(f"Chương {chap_idx}: Khảo Sát Kiến Trúc Hệ Thống Phần {chap_idx}", level=1)
        for sec_idx in range(1, 4):
            doc.add_heading(f"{chap_idx}.{sec_idx} Đánh Giá Hiệu Năng và Khả Năng Mở Rộng", level=2)
            doc.add_paragraph(
                f"Đoạn văn chi tiết cho chương {chap_idx}, mục {sec_idx}. "
                "Hệ thống duy trì trạng thái phân tích tài liệu chuẩn tắc, đảm bảo tính toàn vẹn của mô hình in-memory "
                "khi tài liệu đạt quy mô 50 trang với hàng trăm khối văn bản liên tiếp."
            )
            doc.add_paragraph(
                "Các thao tác gõ phím, chọn vùng, áp dụng giao dịch Agent phải hoạt động với độ trễ tối thiểu "
                "nhờ cơ chế dirty block và cập nhật cục bộ thay vì render lại toàn bộ tài liệu."
            )

    doc.save(str(p))
    return p


def generate_all_fixtures():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    create_simple_docx()
    create_vietnamese_docx()
    create_tables_docx()
    create_complex_formatting_docx()
    create_multi_page_docx()
    create_academic_report_docx()
    create_images_docx()
    create_unsupported_ooxml_docx()
    create_corrupted_docx()
    create_stress_50_page_docx()
    print(f"Generated all fixtures in: {FIXTURES_DIR}")


if __name__ == "__main__":
    generate_all_fixtures()


