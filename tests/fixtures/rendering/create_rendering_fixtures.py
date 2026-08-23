"""
Generator for the 13 Comprehensive Rendering Regression Fixtures:
tests/fixtures/rendering/
- simple.docx
- vietnamese.docx
- headings.docx
- styles.docx
- numbering.docx
- tables.docx
- merged_tables.docx
- images.docx
- sections.docx
- headers_footers.docx
- complex_formatting.docx
- unsupported_ooxml.docx
- stress_50_pages.docx
"""

import os
import sys
from pathlib import Path

if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
import docx
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FIXTURES_DIR = Path(__file__).parent


def create_all_fixtures():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    # 1. simple.docx
    doc1 = docx.Document()
    doc1.add_heading("Tài Liệu Đơn Giản", level=1)
    doc1.add_paragraph("Đây là đoạn văn bản mở đầu thử nghiệm cho công cụ kiểm tra định dạng.")
    doc1.save(str(FIXTURES_DIR / "simple.docx"))

    # 2. vietnamese.docx
    doc2 = docx.Document()
    doc2.add_heading("Kiểm Thử Tiếng Việt Unicode Chuẩn", level=1)
    doc2.add_paragraph("Các ký tự tiếng Việt có dấu: ă â ê ô ơ ư đ. Viết hoa: À Á Ả Ã Ạ, Ắ Ằ Ẳ Ẵ Ặ, Ế Ề Ể Ễ Ệ.")
    doc2.add_paragraph("Ký tự toán học & biểu tượng kỹ thuật: ∀x ∈ ℝ, ∑(i=1..n), λ, π ≈ 3.14159, ✔ ✕ 📐 🤖.")
    doc2.save(str(FIXTURES_DIR / "vietnamese.docx"))

    # 3. headings.docx
    doc3 = docx.Document()
    doc3.add_heading("Chương 1: Cấp độ Tiêu đề 1", level=1)
    doc3.add_paragraph("Nội dung của chương 1.")
    doc3.add_heading("1.1 Mục Tiêu Đề Cấp 2", level=2)
    doc3.add_paragraph("Chi tiết mục 1.1.")
    doc3.add_heading("1.1.1 Mục Tiêu Đề Cấp 3", level=3)
    doc3.add_paragraph("Chi tiết mục 1.1.1.")
    doc3.save(str(FIXTURES_DIR / "headings.docx"))

    # 4. styles.docx
    doc4 = docx.Document()
    doc4.add_heading("Kiểm Thử Kế Thừa Styles", level=1)
    p_style = doc4.add_paragraph()
    r1 = p_style.add_run("Chữ đậm màu xanh dương. ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(2, 132, 199)
    r2 = p_style.add_run("Chữ nghiêng màu xanh lá. ")
    r2.italic = True
    r2.font.color.rgb = RGBColor(22, 163, 74)
    r3 = p_style.add_run("Chữ gạch chân màu đỏ.")
    r3.underline = True
    r3.font.color.rgb = RGBColor(220, 38, 38)
    doc4.save(str(FIXTURES_DIR / "styles.docx"))

    # 5. numbering.docx
    doc5 = docx.Document()
    doc5.add_heading("Kiểm Thử Danh Sách & Numbering", level=1)
    doc5.add_paragraph("Mục danh sách gạch đầu dòng 1", style="List Bullet")
    doc5.add_paragraph("Mục danh sách gạch đầu dòng 2", style="List Bullet")
    doc5.add_paragraph("Mục danh sách đánh số 1", style="List Number")
    doc5.add_paragraph("Mục danh sách đánh số 2", style="List Number")
    doc5.save(str(FIXTURES_DIR / "numbering.docx"))

    # 6. tables.docx
    doc6 = docx.Document()
    doc6.add_heading("Kiểm Thử Bảng Chuẩn", level=1)
    tbl = doc6.add_table(rows=3, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(3):
        for c in range(3):
            tbl.cell(r, c).text = f"Ô ({r+1}, {c+1})"
    doc6.save(str(FIXTURES_DIR / "tables.docx"))

    # 7. merged_tables.docx
    doc7 = docx.Document()
    doc7.add_heading("Kiểm Thử Bảng Gộp Ô (Merged)", level=1)
    tbl_m = doc7.add_table(rows=3, cols=3)
    # Merge horizontal
    cell_00 = tbl_m.cell(0, 0)
    cell_01 = tbl_m.cell(0, 1)
    cell_00.merge(cell_01)
    cell_00.text = "Gộp Cột 1 và 2"
    tbl_m.cell(0, 2).text = "Cột 3"
    # Merge vertical
    cell_10 = tbl_m.cell(1, 0)
    cell_20 = tbl_m.cell(2, 0)
    cell_10.merge(cell_20)
    cell_10.text = "Gộp Hàng 2 và 3"
    tbl_m.cell(1, 1).text = "Ô (2,2)"
    tbl_m.cell(1, 2).text = "Ô (2,3)"
    tbl_m.cell(2, 1).text = "Ô (3,2)"
    tbl_m.cell(2, 2).text = "Ô (3,3)"
    doc7.save(str(FIXTURES_DIR / "merged_tables.docx"))

    # 8. images.docx
    doc8 = docx.Document()
    doc8.add_heading("Kiểm Thử Hình Ảnh", level=1)
    p_img = doc8.add_paragraph("Hình ảnh thử nghiệm minh họa.")
    doc8.save(str(FIXTURES_DIR / "images.docx"))

    # 9. sections.docx
    doc9 = docx.Document()
    sec1 = doc9.sections[0]
    sec1.top_margin = Cm(2.0)
    doc9.add_heading("Section 1: Trang Dọc", level=1)
    doc9.add_paragraph("Nội dung của Section 1.")
    sec2 = doc9.add_section()
    sec2.page_width = Cm(29.7)
    sec2.page_height = Cm(21.0)
    doc9.add_heading("Section 2: Trang Ngang", level=1)
    doc9.add_paragraph("Nội dung của Section 2 nằm ngang.")
    doc9.save(str(FIXTURES_DIR / "sections.docx"))

    # 10. headers_footers.docx
    doc10 = docx.Document()
    s = doc10.sections[0]
    s.different_first_page_header_footer = True
    s.header.paragraphs[0].text = "Header Trang Thường"
    s.footer.paragraphs[0].text = "Footer Trang Thường"
    doc10.add_heading("Trang Bìa (Không Header)", level=1)
    doc10.add_page_break()
    doc10.add_heading("Trang Nội Dung 2 (Có Header & Footer)", level=1)
    doc10.save(str(FIXTURES_DIR / "headers_footers.docx"))

    # 11. complex_formatting.docx
    doc11 = docx.Document()
    doc11.add_heading("Báo Cáo Nghiên Cứu Hệ Thống", level=1)
    p_lead = doc11.add_paragraph()
    p_lead.paragraph_format.first_line_indent = Cm(1.27)
    r_lead = p_lead.add_run("Tóm tắt: ")
    r_lead.bold = True
    p_lead.add_run("Nghiên cứu này trình bày kiến trúc phân tích cú pháp OpenXML độ trung thực cao.")
    doc11.save(str(FIXTURES_DIR / "complex_formatting.docx"))

    # 12. unsupported_ooxml.docx
    doc12 = docx.Document()
    doc12.add_heading("Tài Liệu Bảo Toàn OOXML Lạ", level=1)
    p_unsup = doc12.add_paragraph("Đoạn văn chứa phần tử mở rộng.")
    custom_elem = parse_xml(r'<w:customFutureTag %s w:val="preserved_value"><w:nestedData/></w:customFutureTag>' % nsdecls('w'))
    p_unsup._p.append(custom_elem)
    doc12.save(str(FIXTURES_DIR / "unsupported_ooxml.docx"))

    # 13. stress_50_pages.docx
    doc13 = docx.Document()
    for i in range(1, 51):
        doc13.add_heading(f"Chương {i}: Thử Nghiệm Tải Lớn Trang {i}", level=1)
        for p_idx in range(5):
            doc13.add_paragraph(f"Đoạn văn bản thứ {p_idx+1} của trang {i}. Chứa các mệnh đề kiểm thử hiệu năng phân trang và giải quyết định dạng.")
        if i < 50:
            doc13.add_page_break()
    doc13.save(str(FIXTURES_DIR / "stress_50_pages.docx"))

    print(f"Đã tạo thành công 13 fixtures tại {FIXTURES_DIR}")


if __name__ == "__main__":
    create_all_fixtures()
