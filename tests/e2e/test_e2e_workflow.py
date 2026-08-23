"""
Final E2E Acceptance Test implementing the 15-step master workflow.
"""

from pathlib import Path
import pytest
from PIL import Image
import docx
from docx_agent.agent import DocumentAgent


@pytest.fixture
def test_image_path(temp_dir) -> Path:
    img_p = temp_dir / "test_chart.png"
    img = Image.new("RGB", (200, 100), color=(73, 109, 137))
    img.save(str(img_p))
    return img_p


def test_15_step_e2e_master_workflow(temp_dir, test_image_path):
    # Setup initial fixture document
    doc_path = temp_dir / "initial_report.docx"
    doc = docx.Document()
    
    doc.add_heading("Báo Cáo Dự Án Ban Đầu", level=0)
    doc.add_heading("1. Giới thiệu tổng quan", level=1)
    
    for i in range(10):
        doc.add_paragraph(f"Đoạn văn giới thiệu số {i+1} với thông tin chi tiết về dự án.")
        
    doc.add_heading("2. Phương pháp nghiên cứu", level=1)
    p_mixed = doc.add_paragraph()
    r1 = p_mixed.add_run("Chúng tôi áp dụng ")
    r2 = p_mixed.add_run("thuật toán tối ưu hóa")
    r2.bold = True
    r3 = p_mixed.add_run(" trên toàn bộ tập dữ liệu.")
    
    doc.add_heading("2.1 Chi tiết thuật toán", level=2)
    doc.add_paragraph("Mô tả chi tiết về các tham số đầu vào.")
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Tham số"
    table.rows[0].cells[1].text = "Giá trị"
    table.rows[1].cells[0].text = "Learning Rate"
    table.rows[1].cells[1].text = "0.001"

    doc.save(str(doc_path))

    # -------------------------------------------------------------
    # 15-STEP MASTER WORKFLOW EXECUTION
    # -------------------------------------------------------------

    # Step 1: Inspect document
    agent = DocumentAgent(doc_path)
    summary = agent.inspect()
    assert summary.paragraphs_count >= 13
    assert summary.tables_count == 1

    # Step 2: Apply academic-vn preset
    preset_rep = agent.apply_preset("academic-vn")
    assert preset_rep["preset_applied"] == "academic-vn"

    # Step 3: Change title
    agent.replace("Báo Cáo Dự Án Ban Đầu", "BÁO CÁO NGHIÊN CỨU PHÁT TRIỂN HỆ THỐNG")

    # Step 4: Format headings
    agent.format_text("heading:1", bold=True, font_name="Times New Roman", font_size_pt=16.0)

    # Step 5: Replace text inside mixed-format paragraph
    # Replace 'thuật toán tối ưu hóa' (which is BOLD) with 'mô hình deep learning'
    n_rep = agent.replace("thuật toán tối ưu hóa", "mô hình deep learning")
    assert n_rep == 1

    # Step 6: Edit table
    agent.edit_cell("tbl_0001_r02_c02", text="0.0005", bold=True, bg_color_hex="FFFFCC")

    # Step 7: Insert image
    img_pid = agent.insert_image(
        image_path=str(test_image_path),
        width_cm=8.0,
        alignment="center",
        caption="Hình 1: Biểu đồ kiến trúc mô hình học máy.",
    )
    assert img_pid is not None

    # Step 8: Caption was added during insert_image

    # Step 9: Add page numbers
    agent.set_page_numbers(format_str="Trang {PAGE} / {NUMPAGES}", alignment="center")

    # Step 10: Create TOC
    agent.insert_toc(title="MỤC LỤC TỔNG HỢP")

    # Step 11: Save transactionally
    final_output = temp_dir / "final_report.docx"
    saved_path = agent.save(output_path=final_output, verify=True)
    assert Path(saved_path).exists()

    # Step 12: Reopen document independently
    reopened = docx.Document(saved_path)
    
    # Step 13: Verify
    agent_reopened = DocumentAgent(saved_path)
    v_report = agent_reopened.verify(expected_font="Times New Roman")
    assert v_report["integrity_passed"] is True

    # Check format preservation of step 5
    found_preserved_run = False
    for p in reopened.paragraphs:
        if "mô hình deep learning" in p.text:
            for r in p.runs:
                if "mô hình deep learning" in r.text:
                    assert r.bold is True
                    found_preserved_run = True
    assert found_preserved_run, "Run formatting was not preserved across surgery!"

    # Step 14: Generate diff
    diff_report = DocumentAgent(doc_path).diff(final_output)
    assert not diff_report.identical
    assert diff_report.summary["modified"] > 0

    # Step 15: Return machine-readable report
    machine_report = {
        "success": True,
        "document": str(final_output),
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "diff_summary": diff_report.summary,
        "verification": v_report,
    }
    assert machine_report["success"] is True
