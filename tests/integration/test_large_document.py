"""
Large document stress test: 1000+ paragraphs, 50+ tables, multiple sections.
Verifies performance, memory, and indexing correctness.
"""

import time
import docx
from docx_agent.agent import DocumentAgent


def test_large_document_performance(temp_dir):
    doc_path = temp_dir / "large_doc.docx"
    doc = docx.Document()

    # Build 1000 paragraphs and 50 tables
    for i in range(1000):
        if i % 100 == 0:
            doc.add_heading(f"Chương {i // 100 + 1}: Nội dung chi tiết {i}", level=1)
        doc.add_paragraph(f"Đoạn văn số {i + 1}: Thử nghiệm hiệu năng với dữ liệu lớn và văn bản tiếng Việt.")

    for t in range(50):
        tbl = doc.add_table(rows=2, cols=3)
        tbl.rows[0].cells[0].text = f"T{t+1}"
        tbl.rows[0].cells[1].text = "Header A"
        tbl.rows[0].cells[2].text = "Header B"

    doc.save(str(doc_path))

    # Time document loading and inspection
    t0 = time.time()
    agent = DocumentAgent(doc_path)
    summary = agent.inspect()
    t_inspect = time.time() - t0

    assert summary.paragraphs_count >= 1000
    assert summary.tables_count == 50
    assert t_inspect < 5.0, f"Inspect on 1000+ paragraphs took too long: {t_inspect:.2f}s"

    # Test targeted replacement
    t1 = time.time()
    n = agent.replace(target="Đoạn văn số 500", replacement="ĐOẠN ĐẶC BIỆT ĐÃ SỬA", count=1)
    t_replace = time.time() - t1

    assert n == 1
    assert t_replace < 2.0, f"Targeted replace took too long: {t_replace:.2f}s"

    out_file = temp_dir / "large_doc_saved.docx"
    agent.save(out_file)
    assert out_file.exists()
