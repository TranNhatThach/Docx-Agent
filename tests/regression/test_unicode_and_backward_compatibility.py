"""
Regression tests for Vietnamese Unicode fidelity, mathematical symbols, typography,
and backward-compatible CLI command behaviors.
"""

import docx
from docx_agent.agent import DocumentAgent


def test_vietnamese_unicode_and_special_symbols(temp_dir):
    agent = DocumentAgent()
    
    test_strings = [
        "Tiếng Việt: Đặng, Nguyễn, Trần, Lê, Phạm, Huỳnh, Hoàng, Phan, Vũ, Võ.",
        "Nguyên âm có dấu: ă â ê ô ơ ư Ă Â Ê Ô Ơ Ư á à ả ã ạ.",
        "Typography marks: “smart quotes”, ‘single quotes’, — em dash, … ellipsis.",
        "Mathematical & Greek symbols: →, ≤, ≥, ∑, ∫, √, α, β, γ, θ, λ, π, Ω.",
    ]
    
    for text in test_strings:
        agent.append(text)

    out_file = temp_dir / "unicode_test.docx"
    agent.save(out_file)

    # Reopen independently and verify exact character equivalence
    reopened = docx.Document(str(out_file))
    for idx, expected in enumerate(test_strings):
        actual = reopened.paragraphs[idx].text
        assert actual == expected, f"Unicode mismatch at line {idx}!\nExpected: {expected}\nActual:   {actual}"


def test_vietnamese_text_replacement(temp_dir):
    agent = DocumentAgent()
    agent.append("Đề tài: Ứng dụng trí tuệ nhân tạo trong xử lý ngôn ngữ tự nhiên.")
    
    # Replace Vietnamese accented substring
    n = agent.replace(target="trí tuệ nhân tạo", replacement="mô hình ngôn ngữ lớn (LLM)")
    assert n == 1

    out_file = temp_dir / "vn_replace.docx"
    agent.save(out_file)

    reopened = docx.Document(str(out_file))
    assert "Đề tài: Ứng dụng mô hình ngôn ngữ lớn (LLM) trong xử lý ngôn ngữ tự nhiên." == reopened.paragraphs[0].text
