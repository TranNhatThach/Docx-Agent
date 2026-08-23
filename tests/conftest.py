"""
Pytest configuration and shared fixtures for docx-agent test suite.
"""

import tempfile
import pytest
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx_agent.agent import DocumentAgent


@pytest.fixture
def temp_dir():
    with tempfile.TemporaryDirectory() as tmp:
        yield Path(tmp)


@pytest.fixture
def sample_docx(temp_dir) -> Path:
    """Creates a standard test DOCX file with headings, paragraphs, and formatting."""
    doc_path = temp_dir / "sample.docx"
    doc = docx.Document()
    
    # Title
    doc.add_heading("Báo Cáo Nghiên Cứu Khoa Học", level=0)
    
    # Heading 1
    doc.add_heading("1. Giới thiệu", level=1)
    doc.add_paragraph("Đây là đoạn văn mở đầu giới thiệu về đề tài nghiên cứu.")
    
    # Heading 2 with mixed formatting paragraph
    doc.add_heading("1.1 Phương pháp", level=2)
    p = doc.add_paragraph()
    r1 = p.add_run("Chúng tôi sử dụng ")
    r2 = p.add_run("phương pháp học máy")
    r2.bold = True
    r3 = p.add_run(" để tối ưu hóa kết quả.")
    
    # Table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Thông số"
    table.rows[0].cells[1].text = "Giá trị"
    table.rows[1].cells[0].text = "Độ chính xác"
    table.rows[1].cells[1].text = "98.5%"

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def mixed_runs_docx(temp_dir) -> Path:
    """Creates a test document with intricate mixed run styling."""
    doc_path = temp_dir / "mixed_runs.docx"
    doc = docx.Document()
    
    p = doc.add_paragraph()
    r1 = p.add_run("Hello ")
    
    r2 = p.add_run("world")
    r2.bold = True
    r2.font.color.rgb = RGBColor(255, 0, 0)
    
    r3 = p.add_run(" today and welcome!")
    r3.italic = True

    # Split across runs: 'hel' in run 1, 'lo world' in run 2
    p2 = doc.add_paragraph()
    r2_1 = p2.add_run("hel")
    r2_2 = p2.add_run("lo universe")
    r2_2.bold = True

    doc.save(str(doc_path))
    return doc_path
