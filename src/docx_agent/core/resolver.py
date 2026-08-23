"""
Polymorphic target resolver for locating elements across the document.
"""

import re
from typing import List, Union, Dict, Any, Optional
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.section import Section
from docx_agent.core.document import DocumentModel
from docx_agent.core.exceptions import ElementNotFoundError, AmbiguousTargetError
from docx_agent.utils.unicode import normalize_unicode, normalize_comparison


class TargetResolver:
    """
    Resolves targets specified as IDs, indices, text strings, regex patterns,
    heading structures, or semantic selectors to concrete document elements.
    """

    def __init__(self, model: DocumentModel):
        self.model = model

    def resolve_paragraphs(
        self,
        target: Union[str, int, Dict[str, Any]],
        single: bool = False,
    ) -> List[Paragraph]:
        """
        Resolves one or more paragraphs matching the target specification.
        If `single=True` and multiple match, raises AmbiguousTargetError.
        """
        matched: List[Paragraph] = []

        if isinstance(target, int):
            # Index integer
            if 0 <= target < len(self.model.doc.paragraphs):
                matched.append(self.model.doc.paragraphs[target])

        elif isinstance(target, dict):
            t_type = target.get("type", "").lower()
            if t_type == "heading":
                level = target.get("level")
                text_query = target.get("text", "")
                for p in self.model.doc.paragraphs:
                    s_name = p.style.name if p.style else ""
                    if "heading" in s_name.lower():
                        if level is not None and str(level) not in s_name:
                            continue
                        if not text_query or normalize_comparison(text_query) in normalize_comparison(p.text):
                            matched.append(p)
            elif t_type == "style":
                style_name = target.get("style", "")
                for p in self.model.doc.paragraphs:
                    if p.style and normalize_comparison(p.style.name) == normalize_comparison(style_name):
                        matched.append(p)
            elif "id" in target:
                p = self.model.identity.get_paragraph(target["id"])
                if p:
                    matched.append(p)

        elif isinstance(target, str):
            target_clean = target.strip()
            
            # Element ID: p_0001
            if re.match(r"^p_\d{4}$", target_clean):
                p = self.model.identity.get_paragraph(target_clean)
                if p:
                    matched.append(p)

            # Index syntax: idx:5 or p[5]
            elif target_clean.startswith("idx:"):
                idx_val = int(target_clean.split(":", 1)[1])
                if 0 <= idx_val < len(self.model.doc.paragraphs):
                    matched.append(self.model.doc.paragraphs[idx_val])
            elif re.match(r"^p\[(\d+)\]$", target_clean):
                idx_val = int(re.match(r"^p\[(\d+)\]$", target_clean).group(1))
                if 0 <= idx_val < len(self.model.doc.paragraphs):
                    matched.append(self.model.doc.paragraphs[idx_val])

            # Range syntax: range:p_0001..p_0005 or p_0001..p_0005
            elif ".." in target_clean:
                parts = target_clean.replace("range:", "").split("..")
                if len(parts) == 2:
                    p1_id, p2_id = parts[0].strip(), parts[1].strip()
                    idx1 = self.model.identity.p_idx_by_id.get(p1_id)
                    idx2 = self.model.identity.p_idx_by_id.get(p2_id)
                    if idx1 is not None and idx2 is not None:
                        start_i, end_i = min(idx1, idx2), max(idx1, idx2)
                        matched.extend(self.model.doc.paragraphs[start_i : end_i + 1])

            # Regex syntax: regex:pattern
            elif target_clean.startswith("regex:"):
                pat = target_clean.split(":", 1)[1]
                compiled = re.compile(pat, re.UNICODE)
                for p in self.model.doc.paragraphs:
                    if compiled.search(p.text):
                        matched.append(p)

            # Heading syntax: heading:level:text or heading:text
            elif target_clean.startswith("heading:"):
                parts = target_clean.split(":", 2)
                level_filter = None
                text_filter = ""
                if len(parts) == 3 and parts[1].isdigit():
                    level_filter = parts[1]
                    text_filter = parts[2]
                elif len(parts) == 2:
                    text_filter = parts[1]
                for p in self.model.doc.paragraphs:
                    s_name = p.style.name if p.style else ""
                    if "heading" in s_name.lower():
                        if level_filter and level_filter not in s_name:
                            continue
                        if not text_filter or normalize_comparison(text_filter) in normalize_comparison(p.text):
                            matched.append(p)

            # Style syntax: style:StyleName
            elif target_clean.startswith("style:"):
                st_name = target_clean.split(":", 1)[1]
                for p in self.model.doc.paragraphs:
                    if p.style and normalize_comparison(p.style.name) == normalize_comparison(st_name):
                        matched.append(p)

            # Default: Substring exact text match
            else:
                q_text = target_clean.replace("text:", "")
                for p in self.model.doc.paragraphs:
                    if q_text in p.text or normalize_comparison(q_text) in normalize_comparison(p.text):
                        matched.append(p)

        if not matched:
            raise ElementNotFoundError(target=str(target), target_type="paragraph")

        if single and len(matched) > 1:
            matched_ids = [self.model.identity.get_paragraph_id(p) or "unknown" for p in matched]
            raise AmbiguousTargetError(
                target=str(target),
                match_count=len(matched),
                matched_ids=matched_ids,
            )

        return matched

    def resolve_table(self, target: Union[str, int]) -> Table:
        """Resolves a single table target."""
        if isinstance(target, int):
            if 0 <= target < len(self.model.doc.tables):
                return self.model.doc.tables[target]
        elif isinstance(target, str):
            target_clean = target.strip()
            if re.match(r"^tbl_\d{4}$", target_clean):
                tbl = self.model.identity.get_table(target_clean)
                if tbl:
                    return tbl
            elif target_clean.startswith("idx:"):
                idx_val = int(target_clean.split(":", 1)[1])
                if 0 <= idx_val < len(self.model.doc.tables):
                    return self.model.doc.tables[idx_val]

        raise ElementNotFoundError(target=str(target), target_type="table")

    def resolve_cell(self, target_id: str) -> _Cell:
        """Resolves a cell target (e.g. 'tbl_0001_r01_c02')."""
        cell = self.model.identity.get_cell(target_id)
        if not cell:
            raise ElementNotFoundError(target=target_id, target_type="cell")
        return cell

    def resolve_section(self, target: Union[str, int]) -> Section:
        """Resolves a section target."""
        if isinstance(target, int):
            if 0 <= target < len(self.model.doc.sections):
                return self.model.doc.sections[target]
        elif isinstance(target, str):
            target_clean = target.strip()
            if re.match(r"^sec_\d{4}$", target_clean):
                sec = self.model.identity.get_section(target_clean)
                if sec:
                    return sec
            elif target_clean.startswith("idx:"):
                idx_val = int(target_clean.split(":", 1)[1])
                if 0 <= idx_val < len(self.model.doc.sections):
                    return self.model.doc.sections[idx_val]

        raise ElementNotFoundError(target=str(target), target_type="section")
