"""
DocumentModel: Object-oriented abstraction wrapping python-docx with agent-centric capabilities.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_agent.core.elements import (
    ElementType,
    ParagraphInfo,
    RunInfo,
    TableInfo,
    SectionInfo,
    DocumentSummary,
)
from docx_agent.core.identity import IdentityManager
from docx_agent.core.exceptions import DocumentNotFoundError, InvalidDocxError
from docx_agent.utils.paths import resolve_safe_path
from docx_agent.utils.unicode import normalize_unicode


from docx.document import Document as DocxDocument


class DocumentModel:
    """
    High-level document model encapsulating OOXML hierarchy, deterministic identities,
    and structural inspection for AI agents.
    """

    def __init__(self, doc_or_path: Union[str, Path, DocxDocument]):
        if isinstance(doc_or_path, (str, Path)):
            path = resolve_safe_path(doc_or_path)
            if not path.exists():
                raise DocumentNotFoundError(str(path))
            try:
                self.doc = docx.Document(str(path))
            except Exception as e:
                raise InvalidDocxError(str(path), str(e))
            self.file_path: Optional[str] = str(path)
        elif isinstance(doc_or_path, DocxDocument):
            self.doc = doc_or_path
            self.file_path = None
        else:
            raise TypeError("doc_or_path must be a file path or DocxDocument instance")

        self.identity = IdentityManager(self.doc)

    def reload(self) -> None:
        """Reloads the document from file if a path is associated."""
        if self.file_path:
            self.doc = docx.Document(self.file_path)
            self.identity = IdentityManager(self.doc)

    def reindex(self) -> None:
        """Refreshes identity registry after mutations."""
        self.identity.reindex()

    def get_summary(self) -> DocumentSummary:
        """Generates a comprehensive summary of the document structure."""
        styles_set = set()
        for p in self.doc.paragraphs:
            if p.style:
                styles_set.add(p.style.name)
        for tbl in self.doc.tables:
            if tbl.style:
                styles_set.add(tbl.style.name)

        outline = self.get_outline()
        sections_info = self.get_sections()

        return DocumentSummary(
            file_path=self.file_path,
            paragraphs_count=len(self.doc.paragraphs),
            tables_count=len(self.doc.tables),
            sections_count=len(self.doc.sections),
            styles=sorted(list(styles_set)),
            headings_outline=outline,
            sections=sections_info,
        )

    def get_paragraphs(
        self,
        start_idx: int = 0,
        end_idx: Optional[int] = None,
        include_runs: bool = False,
    ) -> List[ParagraphInfo]:
        """Returns structured paragraph metadata within index range."""
        result: List[ParagraphInfo] = []
        end = end_idx if end_idx is not None else len(self.doc.paragraphs)
        end = min(end, len(self.doc.paragraphs))

        for idx in range(start_idx, end):
            pid = f"p_{idx + 1:04d}"
            p = self.doc.paragraphs[idx]
            text = normalize_unicode(p.text)
            style_name = p.style.name if p.style else "Normal"
            is_heading = style_name.startswith("Heading") or style_name.lower().startswith("heading")
            heading_level = None
            if is_heading:
                try:
                    num_str = "".join(filter(str.isdigit, style_name))
                    heading_level = int(num_str) if num_str else 1
                except ValueError:
                    heading_level = 1

            align_str = None
            if p.alignment is not None:
                align_str = str(p.alignment).split(".")[-1].lower()

            runs_list: Optional[List[RunInfo]] = None
            if include_runs:
                runs_list = []
                for r_idx, r in enumerate(p.runs):
                    font_size = r.font.size.pt if r.font and r.font.size else None
                    font_name = r.font.name if r.font else None
                    color_hex = str(r.font.color.rgb) if r.font and r.font.color and r.font.color.rgb else None
                    runs_list.append(
                        RunInfo(
                            id=f"{pid}_r{r_idx + 1:02d}",
                            text=normalize_unicode(r.text),
                            font_name=font_name,
                            font_size_pt=font_size,
                            bold=r.bold,
                            italic=r.italic,
                            underline=bool(r.underline),
                            strike=r.font.strike if r.font else None,
                            color_rgb=color_hex,
                            superscript=r.font.superscript if r.font else None,
                            subscript=r.font.subscript if r.font else None,
                        )
                    )

            result.append(
                ParagraphInfo(
                    id=pid,
                    type=ElementType.PARAGRAPH,
                    index=idx,
                    text=text,
                    style=style_name,
                    is_heading=is_heading,
                    heading_level=heading_level,
                    alignment=align_str,
                    runs_count=len(p.runs),
                    runs=runs_list,
                )
            )
        return result

    def get_outline(self) -> List[Dict[str, Any]]:
        """Extracts document heading outline."""
        outline = []
        for idx, p in enumerate(self.doc.paragraphs):
            style_name = p.style.name if p.style else ""
            if "Heading" in style_name or "heading" in style_name.lower() or style_name in ["Title", "Subtitle"]:
                level = 1
                if "1" in style_name:
                    level = 1
                elif "2" in style_name:
                    level = 2
                elif "3" in style_name:
                    level = 3
                elif "4" in style_name:
                    level = 4
                elif style_name == "Title":
                    level = 0
                elif style_name == "Subtitle":
                    level = 1
                outline.append(
                    {
                        "id": f"p_{idx + 1:04d}",
                        "index": idx,
                        "style": style_name,
                        "level": level,
                        "text": normalize_unicode(p.text).strip(),
                    }
                )
        return outline

    def get_tables(self) -> List[TableInfo]:
        """Returns structured metadata for all tables in document."""
        tables_info = []
        for idx, tbl in enumerate(self.doc.tables):
            tid = f"tbl_{idx + 1:04d}"
            preview = []
            for r in tbl.rows[:5]:  # Preview up to 5 rows
                row_cells = [normalize_unicode(c.text).strip().replace("\n", " ") for c in r.cells]
                preview.append(row_cells)

            tables_info.append(
                TableInfo(
                    id=tid,
                    type=ElementType.TABLE,
                    index=idx,
                    rows=len(tbl.rows),
                    columns=len(tbl.columns) if tbl.rows else 0,
                    style=tbl.style.name if tbl.style else None,
                    preview=preview,
                )
            )
        return tables_info

    def get_sections(self) -> List[SectionInfo]:
        """Returns metadata for all sections."""
        secs = []
        for idx, sec in enumerate(self.doc.sections):
            sid = f"sec_{idx + 1:04d}"
            orient = "portrait"
            if sec.page_width > sec.page_height:
                orient = "landscape"
            secs.append(
                SectionInfo(
                    id=sid,
                    type=ElementType.SECTION,
                    index=idx,
                    page_width_cm=round(sec.page_width.cm, 2),
                    page_height_cm=round(sec.page_height.cm, 2),
                    orientation=orient,
                    margin_top_cm=round(sec.top_margin.cm, 2),
                    margin_bottom_cm=round(sec.bottom_margin.cm, 2),
                    margin_left_cm=round(sec.left_margin.cm, 2),
                    margin_right_cm=round(sec.right_margin.cm, 2),
                    has_header=bool(sec.header and sec.header.paragraphs and any(p.text for p in sec.header.paragraphs)),
                    has_footer=bool(sec.footer and sec.footer.paragraphs and any(p.text for p in sec.footer.paragraphs)),
                )
            )
        return secs

    def get_capabilities_report(self) -> Dict[str, Any]:
        """Reports document capabilities, supported operations, and any potential unsupported elements."""
        supported = [
            "read", "inspect", "outline", "find", "replace", "insert", "delete",
            "format_text", "format_paragraph", "styles", "tables", "images",
            "sections", "headers", "footers", "fields", "toc", "presets", "diff", "verify"
        ]
        unsupported = []
        warnings = []

        # Check for drawing/vml/shapes or embedded objects
        xml_str = self.doc._body._element.xml
        if "w:drawing" in xml_str:
            supported.append("drawings_detected")
        if "w:object" in xml_str or "w:embeddedObject" in xml_str:
            unsupported.append("embedded_ole_objects")
            warnings.append("Document contains embedded OLE objects which cannot be edited directly.")
        if "w:txbxContent" in xml_str:
            warnings.append("Document contains floating text boxes (w:txbxContent). Text in boxes is not in main paragraph flow.")

        return {
            "file": self.file_path,
            "supported_operations": supported,
            "unsupported_features": unsupported,
            "warnings": warnings,
            "can_safely_edit": len(unsupported) == 0,
        }
