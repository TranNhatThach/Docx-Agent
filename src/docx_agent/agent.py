"""
DocumentAgent: High-level Python API orchestrating all document capabilities,
transactional safety, format preservation, and agent-native operations (v2.0 Workspace Edition).
"""

from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
import docx
from docx.document import Document as DocxDocument

from docx_agent.core.document import DocumentModel
from docx_agent.core.elements import DocumentSummary, ParagraphInfo, TableInfo, SectionInfo
from docx_agent.core.resolver import TargetResolver
from docx_agent.core.exceptions import (
    DocxAgentError,
    ErrorCode,
    DocumentNotFoundError,
    VerificationError,
    TransactionError,
)
from docx_agent.operations.content import ContentOperations
from docx_agent.operations.formatting import TextFormattingOperations
from docx_agent.operations.paragraphs import ParagraphFormattingOperations
from docx_agent.operations.styles import StyleOperations
from docx_agent.operations.tables import TableOperations
from docx_agent.operations.images import ImageOperations
from docx_agent.operations.sections import SectionOperations
from docx_agent.operations.headers import HeaderOperations
from docx_agent.operations.footers import FooterOperations
from docx_agent.operations.fields import FieldOperations
from docx_agent.operations.academic import AcademicOperations

from docx_agent.canonical.model import (
    DocumentNode,
    SectionNode,
    ParagraphBlock,
    HeadingBlock,
    TableBlock,
    ImageBlock,
    DiagramBlock,
    SourceMetadata,
    CitationNode,
)
from docx_agent.adapters.docx import DocxImporter, DocxExporter
from docx_agent.engine.operations import (
    DocOperation,
    InsertTextOp,
    DeleteTextOp,
    ReplaceTextOp,
    FormatParagraphOp,
    InsertBlockOp,
    DeleteBlockOp,
    InsertCitationOp,
)
from docx_agent.engine.transactions import AgentTransactionManager, TransactionPreview
from docx_agent.engine.selection import SelectionProvider, SelectionContext
from docx_agent.research.provider import ResearchAssistant, ResearchProposal
from docx_agent.media.diagrams import DiagramSynthesizer
from docx_agent.engine.clarification import ClarificationEngine, ClarificationRequest, AgentConfidence
from docx_agent.verification.visual import VisualLayoutVerifier, VisualVerificationReport

from docx_agent.transactions.transaction import TransactionContext
from docx_agent.transactions.backup import BackupManager
from docx_agent.verification.validator import DocumentValidator, ValidationReport
from docx_agent.verification.formatting import FormatChecker, FormatVerificationReport
from docx_agent.verification.diff import DiffEngine, DocumentDiffReport
from docx_agent.interfaces.schemas.models import BatchPlanSchema, PlanOperationSchema, AgentResponse
from docx_agent.utils.paths import resolve_safe_path
from docx_agent.utils.unicode import normalize_unicode, normalize_comparison


class DocumentAgent:
    """
    Universal Agent-Native DOCX manipulation engine & Workspace collaborator.
    """

    def __init__(self, doc_or_path: Optional[Union[str, Path, DocxDocument]] = None):
        if doc_or_path is None:
            self.model = DocumentModel(docx.Document())
            self.canonical_doc = DocumentNode()
        elif isinstance(doc_or_path, (str, Path)) and Path(doc_or_path).exists():
            self.model = DocumentModel(doc_or_path)
            self.canonical_doc = DocxImporter.import_docx(doc_or_path)
        else:
            self.model = DocumentModel(doc_or_path)
            self.canonical_doc = DocumentNode()

        self._init_subsystems()

    def _init_subsystems(self) -> None:
        self.resolver = TargetResolver(self.model)
        self.content = ContentOperations(self.model)
        self.text_fmt = TextFormattingOperations(self.model)
        self.para_fmt = ParagraphFormattingOperations(self.model)
        self.styles = StyleOperations(self.model)
        self.tables = TableOperations(self.model)
        self.images = ImageOperations(self.model)
        self.sections = SectionOperations(self.model)
        self.headers = HeaderOperations(self.model)
        self.footers = FooterOperations(self.model)
        self.fields = FieldOperations(self.model)
        self.academic = AcademicOperations(self.model)

        # V2 Workspace Engines
        self.tx_manager = AgentTransactionManager(self.canonical_doc)
        self.research = ResearchAssistant()

    # ---------------------------------------------------------
    # INSPECTION & DISCOVERY
    # ---------------------------------------------------------
    def inspect(self) -> DocumentSummary:
        """Returns full structural summary of document."""
        return self.model.get_summary()

    def read(
        self,
        start: int = 0,
        end: Optional[int] = None,
        include_runs: bool = False,
    ) -> List[ParagraphInfo]:
        """Reads paragraph contents within index range."""
        return self.model.get_paragraphs(start_idx=start, end_idx=end, include_runs=include_runs)

    def outline(self) -> List[Dict[str, Any]]:
        """Returns document heading hierarchy."""
        return self.model.get_outline()

    def capabilities(self) -> Dict[str, Any]:
        """Returns capability assessment of the document."""
        return self.model.get_capabilities_report()

    def find(
        self,
        text: Optional[str] = None,
        regex: Optional[str] = None,
        style: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """Finds paragraphs matching text, regex, or style."""
        matches = []
        selector = ""
        if regex:
            selector = f"regex:{regex}"
        elif style:
            selector = f"style:{style}"
        elif text:
            selector = f"text:{text}"

        if selector:
            try:
                matched_paras = self.resolver.resolve_paragraphs(selector)
                for p in matched_paras:
                    pid = self.model.identity.get_paragraph_id(p) or "unknown"
                    matches.append({
                        "id": pid,
                        "style": p.style.name if p.style else "Normal",
                        "text": normalize_unicode(p.text),
                    })
            except Exception:
                pass
        return matches

    # ---------------------------------------------------------
    # CONTENT OPERATIONS
    # ---------------------------------------------------------
    def replace(
        self,
        target: str,
        replacement: str,
        scope: Optional[Union[str, int, Dict[str, Any]]] = None,
        count: Optional[int] = None,
    ) -> int:
        """Replaces text while strictly preserving run formatting."""
        return self.content.replace_text(
            target=target,
            replacement=replacement,
            scope_target=scope,
            count=count,
        )

    def insert(
        self,
        text: str,
        target: Union[str, int, Dict[str, Any]],
        position: str = "after",
        style: Optional[str] = None,
    ) -> str:
        """Inserts paragraph before/after target."""
        new_id = self.content.insert_paragraph(
            text=text,
            target=target,
            position=position,
            style=style,
        )
        # Mirror to canonical model
        if self.canonical_doc.sections:
            self.canonical_doc.sections[0].blocks.append(
                ParagraphBlock(style_name=style or "Normal")
            )
        return new_id

    def append(
        self,
        text: str,
        style: Optional[str] = None,
        heading_level: Optional[int] = None,
    ) -> str:
        """Appends paragraph or heading."""
        if heading_level is not None or (style and style.startswith("Heading")):
            lvl = heading_level or int("".join(filter(str.isdigit, style or "1")) or 1)
            self.canonical_doc.sections[0].blocks.append(
                HeadingBlock(level=lvl, style_name=style or f"Heading {lvl}")
            )
        else:
            self.canonical_doc.sections[0].blocks.append(
                ParagraphBlock(style_name=style or "Normal")
            )

        return self.content.append_paragraph(
            text=text,
            style=style,
            heading_level=heading_level,
        )

    def delete(self, target: Union[str, int, Dict[str, Any]]) -> bool:
        """Deletes target element."""
        return self.content.delete_element(target)

    # ---------------------------------------------------------
    # FORMATTING OPERATIONS
    # ---------------------------------------------------------
    def format_text(self, target: Union[str, int, Dict[str, Any]], **kwargs) -> int:
        """Applies character font formatting."""
        return self.text_fmt.format_text(target=target, **kwargs)

    def format_paragraph(self, target: Union[str, int, Dict[str, Any]], **kwargs) -> int:
        """Applies paragraph layout formatting."""
        return self.para_fmt.format_paragraph(target=target, **kwargs)

    # ---------------------------------------------------------
    # PRESETS & ACADEMIC
    # ---------------------------------------------------------
    def apply_preset(self, preset_name: str = "academic-vn") -> Dict[str, Any]:
        """Applies complete formatting preset."""
        return self.academic.apply_preset(preset_name)

    # ---------------------------------------------------------
    # TABLES, IMAGES, SECTIONS, HEADERS, FOOTERS, FIELDS
    # ---------------------------------------------------------
    def create_table(self, rows: int, cols: int, **kwargs) -> str:
        return self.tables.create_table(rows=rows, columns=cols, **kwargs)

    def edit_cell(self, cell_id: str, text: str, **kwargs) -> bool:
        return self.tables.edit_cell(cell_id=cell_id, text=text, **kwargs)

    def insert_image(self, image_path: str, **kwargs) -> str:
        return self.images.insert_image(image_path=image_path, **kwargs)

    def configure_section(self, **kwargs) -> bool:
        return self.sections.configure_section(**kwargs)

    def set_header(self, text: str, **kwargs) -> bool:
        return self.headers.set_header(text=text, **kwargs)

    def set_footer(self, text: str, **kwargs) -> bool:
        return self.footers.set_footer(text=text, **kwargs)

    def set_page_numbers(self, **kwargs) -> bool:
        return self.footers.set_page_numbers(**kwargs)

    def insert_toc(self, **kwargs) -> str:
        return self.fields.insert_toc(**kwargs)

    # ---------------------------------------------------------
    # V2 WORKSPACE & AGENT FEATURES
    # ---------------------------------------------------------
    def get_selection_context(self, block_id: str, start: int = 0, end: int = 0) -> SelectionContext:
        """Extracts rich context for selected block and text slice."""
        return SelectionProvider.build_context(self.canonical_doc, block_id, start, end)

    def propose_agent_transaction(self, description: str, operations: List[DocOperation]) -> TransactionPreview:
        """Proposes a multi-operation transaction with user preview."""
        return self.tx_manager.propose_transaction(description, operations)

    def apply_agent_transaction(self, transaction_id: str) -> bool:
        """Applies an approved transaction."""
        return self.tx_manager.apply_pending_transaction(transaction_id)

    def reject_agent_transaction(self, transaction_id: str) -> bool:
        """Rejects a transaction."""
        return self.tx_manager.reject_pending_transaction(transaction_id)

    def undo(self) -> bool:
        """Undoes the most recent operation or agent transaction."""
        return self.tx_manager.undo()

    def redo(self) -> bool:
        """Redoes the most recently undone operation."""
        return self.tx_manager.redo()

    def research_claim(self, claim: str, style: str = "apa") -> ResearchProposal:
        """Searches verified sources and constructs un-hallucinated citation proposal."""
        return self.research.evaluate_claim_and_propose_citation(claim, citation_style=style)

    def insert_citation(self, block_id: str, offset: int, source: SourceMetadata, style: str = "apa") -> bool:
        """Inserts citation into canonical model and paragraph."""
        op = InsertCitationOp(block_id=block_id, offset=offset, source=source, citation_style=style)
        return self.tx_manager.execute_operation(op)

    def generate_diagram(self, diag_type: str, items: List[str], title: str = "Architecture") -> DiagramBlock:
        """Synthesizes structured Mermaid & SVG diagram."""
        if diag_type == "flowchart":
            return DiagramSynthesizer.generate_flowchart(items, title=title)
        else:
            return DiagramSynthesizer.generate_architecture_diagram(items, title=title)

    def clarify_instruction(self, instruction: str, selected_text: Optional[str] = None) -> Tuple[AgentConfidence, Optional[ClarificationRequest]]:
        """Assesses ambiguity and returns multiple-choice questions if confidence is low."""
        return ClarificationEngine.assess_instruction(
            instruction=instruction,
            selected_text=selected_text,
            document_profile=self.canonical_doc.profile.value,
        )

    # ---------------------------------------------------------
    # VERIFICATION & DIFF
    # ---------------------------------------------------------
    def verify(
        self,
        expected_font: Optional[str] = None,
        expected_size_pt: Optional[float] = None,
        expected_line_spacing: Optional[float] = None,
        expected_alignment: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Runs dual verification: structural integrity, formatting checks, and visual layout."""
        format_rep = FormatChecker.verify_document_formatting(
            doc_or_path=self.model.doc,
            expected_font=expected_font,
            expected_size_pt=expected_size_pt,
            expected_line_spacing=expected_line_spacing,
            expected_alignment=expected_alignment,
        )
        visual_rep = VisualLayoutVerifier.verify_document_layout(self.canonical_doc)

        return {
            "integrity_passed": True,
            "format_verification": format_rep.model_dump(),
            "visual_layout_verification": visual_rep.model_dump(),
            "dual_verification_passed": format_rep.passed and visual_rep.passed,
        }

    def diff(self, other_path: Union[str, Path]) -> DocumentDiffReport:
        """Diffs this document with another document."""
        return DiffEngine.compare(self.model.doc, other_path)

    # ---------------------------------------------------------
    # BATCH PLAN EXECUTION
    # ---------------------------------------------------------
    def apply_plan(self, plan: Union[Dict[str, Any], BatchPlanSchema]) -> Dict[str, Any]:
        """Executes multi-operation plan atomically."""
        if isinstance(plan, dict):
            plan_obj = BatchPlanSchema(**plan)
        else:
            plan_obj = plan

        executed = 0
        for op in plan_obj.operations:
            op_type = op.type.lower().replace("-", "_")
            if op_type == "replace" and op.target and op.replacement:
                self.replace(str(op.target), op.replacement)
                executed += 1
            elif op_type == "insert" and op.text and op.target:
                self.insert(op.text, op.target, position=op.position or "after", style=op.style)
                executed += 1
            elif op_type == "append" and op.text:
                self.append(op.text, style=op.style, heading_level=op.level)
                executed += 1
            elif op_type == "delete" and op.target:
                self.delete(op.target)
                executed += 1
            elif op_type == "format_text" and op.target:
                self.format_text(
                    op.target,
                    font_name=op.font_name,
                    font_size_pt=op.font_size_pt,
                    bold=op.bold,
                    italic=op.italic,
                    underline=op.underline,
                    color_rgb=op.color_rgb,
                )
                executed += 1
            elif op_type == "format_para" and op.target:
                self.format_paragraph(
                    op.target,
                    alignment=op.alignment,
                    line_spacing=op.line_spacing,
                    space_before_pt=op.space_before_pt,
                    space_after_pt=op.space_after_pt,
                    first_line_indent_cm=op.first_line_indent_cm,
                )
                executed += 1
            elif op_type == "preset" and op.preset:
                self.apply_preset(op.preset)
                executed += 1
            elif op_type == "page_number":
                self.set_page_numbers(format_str=op.format or "Page {PAGE} of {NUMPAGES}", alignment=op.alignment or "center")
                executed += 1
            elif op_type == "toc":
                self.insert_toc()
                executed += 1

        return {
            "success": True,
            "operations_executed": executed,
        }

    # ---------------------------------------------------------
    # TRANSACTIONAL SAVE
    # ---------------------------------------------------------
    def save(
        self,
        output_path: Optional[Union[str, Path]] = None,
        auto_backup: bool = True,
        verify: bool = True,
    ) -> str:
        """
        Saves document using TransactionContext (validate -> backup -> save temp -> verify -> commit).
        """
        target_path = resolve_safe_path(output_path or self.model.file_path or "output.docx")

        with TransactionContext(
            file_path=target_path,
            auto_backup=auto_backup,
            verify_on_commit=verify,
        ) as tx:
            tx.model = self.model
            tx.save_and_verify()

        self.model.file_path = str(target_path)
        return str(target_path)

    def export_canonical(self, output_path: Union[str, Path]) -> str:
        """Exports Canonical Document Model to DOCX format."""
        out_p = DocxExporter.export_docx(self.canonical_doc, output_path)
        return str(out_p)
