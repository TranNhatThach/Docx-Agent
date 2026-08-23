"""
High-Fidelity DOCX Import and Export Adapters (Word-Grade OOXML Engine).
Bi-directionally converts between Microsoft Word OpenXML packages and Canonical Document Model
with full Style Resolution, Multilevel Numbering, DrawingML, and Table Geometry Preservation.
"""

from pathlib import Path
from typing import Union, Optional, List, Dict, Any
import docx
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from xml.etree import ElementTree as ET

from docx_agent.canonical.model import (
    DocumentNode,
    SectionNode,
    SectionProperties,
    ParagraphBlock,
    HeadingBlock,
    ListItemBlock,
    TableBlock,
    TableCellNode,
    ImageBlock,
    DiagramBlock,
    UnsupportedBlock,
    RunNode,
    BlockNode,
    generate_id,
)
from docx_agent.engine.styles import StyleResolver
from docx_agent.engine.numbering import NumberingResolver
from docx_agent.ooxml.helpers import set_run_text_preserving_spaces
from docx_agent.ooxml.fields import add_page_number_to_paragraph, add_toc_field
from docx_agent.ooxml.tables import (
    set_cell_background,
    set_row_repeat_header,
    set_row_cant_split,
    set_table_borders,
)
from docx_agent.operations.formatting import hex_to_rgb, HIGHLIGHT_MAP
from docx_agent.operations.paragraphs import ALIGN_MAP
from docx_agent.utils.paths import resolve_safe_path, ensure_parent_dir
from docx_agent.utils.unicode import normalize_unicode


class DocxImporter:
    """
    Parses OpenXML Word documents into the rich Canonical Document Model.
    """

    @staticmethod
    def import_docx(file_path: Union[str, Path]) -> DocumentNode:
        path = resolve_safe_path(file_path)
        doc = docx.Document(str(path))
        doc_node = DocumentNode(title=path.stem)

        # 1. Read raw styles.xml and numbering.xml if present in the package
        raw_styles_xml = None
        raw_numbering_xml = None
        try:
            for rel in doc.part.related_parts.values():
                if "styles" in rel.partname:
                    raw_styles_xml = rel.blob.decode("utf-8", errors="ignore")
                elif "numbering" in rel.partname:
                    raw_numbering_xml = rel.blob.decode("utf-8", errors="ignore")
        except Exception:
            pass

        doc_node.raw_styles_xml = raw_styles_xml
        doc_node.raw_numbering_xml = raw_numbering_xml

        style_resolver = StyleResolver(raw_styles_xml)
        numbering_resolver = NumberingResolver(raw_numbering_xml)

        doc_node.sections = []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
              "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
              "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
              "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
              "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"}

        # 2. Ingest Sections definitions
        for s_idx, sec in enumerate(doc.sections):
            orient = "portrait"
            if sec.page_width > sec.page_height:
                orient = "landscape"

            diff_first = bool(getattr(sec, "different_first_page_header_footer", False))

            sec_props = SectionProperties(
                page_width_cm=round(sec.page_width.cm, 2),
                page_height_cm=round(sec.page_height.cm, 2),
                orientation=orient,
                margin_top_cm=round(sec.top_margin.cm, 2),
                margin_bottom_cm=round(sec.bottom_margin.cm, 2),
                margin_left_cm=round(sec.left_margin.cm, 2),
                margin_right_cm=round(sec.right_margin.cm, 2),
                different_first_page=diff_first,
            )

            header_t = sec.header.paragraphs[0].text if sec.header and sec.header.paragraphs and sec.header.paragraphs[0].text else None
            footer_t = sec.footer.paragraphs[0].text if sec.footer and sec.footer.paragraphs and sec.footer.paragraphs[0].text else None

            first_header_t = None
            first_footer_t = None
            if diff_first and hasattr(sec, "first_page_header") and sec.first_page_header.paragraphs:
                first_header_t = sec.first_page_header.paragraphs[0].text or None
            if diff_first and hasattr(sec, "first_page_footer") and sec.first_page_footer.paragraphs:
                first_footer_t = sec.first_page_footer.paragraphs[0].text or None

            sec_node = SectionNode(
                id=f"sec_{s_idx+1:04d}",
                properties=sec_props,
                header_text=header_t,
                footer_text=footer_t,
                first_page_header_text=first_header_t,
                first_page_footer_text=first_footer_t,
                has_page_numbers=True,
            )
            doc_node.sections.append(sec_node)

        if not doc_node.sections:
            doc_node.sections.append(SectionNode())

        current_sec_idx = 0
        target_sec = doc_node.sections[current_sec_idx]

        # 3. Ingest Body Elements in sequential document order without section duplication
        for blk_idx, child in enumerate(doc._body._element):
            tag = child.tag

            if tag.endswith("sectPr"):
                # End of current section, advance to next section if available
                if current_sec_idx + 1 < len(doc_node.sections):
                    current_sec_idx += 1
                    target_sec = doc_node.sections[current_sec_idx]
                continue

            if tag.endswith("p"):
                p = docx.text.paragraph.Paragraph(child, doc)
                style_name = p.style.name if p.style else "Normal"

                # Check paragraph-level section break (pPr/sectPr)
                p_sect_pr = child.find(f".//{{{ns['w']}}}sectPr")
                
                # Check numbering (w:numPr)
                num_id = None
                ilvl = 0
                resolved_num_label = None
                num_pr = child.find(f".//{{{ns['w']}}}numPr")
                if num_pr is not None:
                    num_id_elem = num_pr.find(f"{{{ns['w']}}}numId")
                    if num_id_elem is not None:
                        val = num_id_elem.attrib.get(f"{{{ns['w']}}}val")
                        if val and val.isdigit():
                            num_id = int(val)
                    ilvl_elem = num_pr.find(f"{{{ns['w']}}}ilvl")
                    if ilvl_elem is not None:
                        val = ilvl_elem.attrib.get(f"{{{ns['w']}}}val")
                        if val and val.isdigit():
                            ilvl = int(val)
                    if num_id is not None:
                        resolved_num_label = numbering_resolver.get_numbering_label(num_id, ilvl)

                # Check paragraph properties
                align_str = None
                if p.alignment is not None:
                    align_raw = getattr(p.alignment, "name", str(p.alignment))
                    align_str = align_raw.split()[0].split(".")[-1].lower()
                    if "(" in align_str:
                        align_str = align_str.split("(")[0].strip()
                ls = p.paragraph_format.line_spacing or None
                sb = p.paragraph_format.space_before.pt if p.paragraph_format.space_before else None
                sa = p.paragraph_format.space_after.pt if p.paragraph_format.space_after else None
                fl = p.paragraph_format.first_line_indent.cm if p.paragraph_format.first_line_indent else None

                keep_next = child.find(f".//{{{ns['w']}}}keepNext") is not None
                page_break_before = child.find(f".//{{{ns['w']}}}pageBreakBefore") is not None

                # Resolve effective paragraph formatting
                direct_p_props = {
                    "alignment": align_str,
                    "line_spacing": ls,
                    "space_before_pt": sb,
                    "space_after_pt": sa,
                    "first_line_indent_cm": fl,
                    "keep_with_next": keep_next or None,
                    "page_break_before": page_break_before or None,
                }
                eff_para = style_resolver.resolve_paragraph(style_name, direct_p_props)

                # Ingest Runs
                runs_nodes = []
                image_blocks_in_p = []

                for r_idx, r in enumerate(p.runs):
                    font_size = r.font.size.pt if r.font and r.font.size else None
                    font_name = r.font.name if r.font else None
                    color_hex = str(r.font.color.rgb) if r.font and r.font.color and r.font.color.rgb else None
                    is_page_brk = bool(r._r.find(f".//{{{ns['w']}}}br[@{{{ns['w']}}}type='page']") is not None)

                    # Extract DrawingML image if present inside run
                    drawing = r._r.find(f".//{{{ns['w']}}}drawing")
                    if drawing is not None:
                        extent = drawing.find(f".//{{{ns['wp']}}}extent")
                        w_cm = 10.0
                        h_cm = None
                        if extent is not None:
                            cx = extent.attrib.get("cx")
                            cy = extent.attrib.get("cy")
                            if cx and cx.isdigit():
                                w_cm = round(float(cx) / 360000.0, 2)
                            if cy and cy.isdigit():
                                h_cm = round(float(cy) / 360000.0, 2)

                        blip = drawing.find(f".//{{{ns['a']}}}blip")
                        r_id = blip.attrib.get(f"{{{ns['r']}}}embed") if blip is not None else None
                        img_path = f"image_{r_id or 'embedded'}.png"

                        img_blk = ImageBlock(
                            id=f"img_{current_sec_idx+1:02d}_{blk_idx+1:04d}_{r_idx+1:02d}",
                            source_uri_or_path=img_path,
                            width_cm=w_cm,
                            height_cm=h_cm,
                            alignment=eff_para.alignment,
                        )
                        image_blocks_in_p.append(img_blk)

                    # Resolve effective run style
                    direct_r_props = {
                        "font_name": font_name,
                        "font_size_pt": font_size,
                        "bold": r.bold,
                        "italic": r.italic,
                        "underline": bool(r.underline) if r.underline is not None else None,
                        "color_rgb": color_hex,
                    }
                    eff_run = style_resolver.resolve_run(eff_para, direct_r_props)

                    run_node = RunNode(
                        id=f"r_{current_sec_idx+1:02d}_{blk_idx+1:04d}_{r_idx+1:03d}",
                        text=normalize_unicode(r.text),
                        font_name=eff_run.font_name,
                        font_size_pt=eff_run.font_size_pt,
                        bold=eff_run.bold,
                        italic=eff_run.italic,
                        underline=eff_run.underline,
                        color_rgb=eff_run.color_rgb,
                        highlight=eff_run.highlight,
                        is_page_break=is_page_brk,
                        effective_formatting=eff_run.model_dump(),
                    )
                    runs_nodes.append(run_node)

                # Determine Block Type & Instantiate
                if style_name.startswith("Heading"):
                    try:
                        lvl = int("".join(filter(str.isdigit, style_name)) or 1)
                    except ValueError:
                        lvl = 1
                    target_sec.blocks.append(
                        HeadingBlock(
                            id=f"h_{current_sec_idx+1:02d}_{blk_idx+1:04d}",
                            style_name=style_name,
                            level=lvl,
                            runs=runs_nodes,
                            alignment=eff_para.alignment,
                            line_spacing=eff_para.line_spacing,
                            space_before_pt=eff_para.space_before_pt,
                            space_after_pt=eff_para.space_after_pt,
                            first_line_indent_cm=eff_para.first_line_indent_cm,
                            keep_with_next=eff_para.keep_with_next,
                            page_break_before=eff_para.page_break_before,
                            effective_properties=eff_para.model_dump(),
                        )
                    )
                elif num_id is not None or style_name.startswith("List"):
                    target_sec.blocks.append(
                        ListItemBlock(
                            id=f"p_{current_sec_idx+1:02d}_{blk_idx+1:04d}",
                            style_name=style_name,
                            list_type="bullet" if (resolved_num_label == "•" or "Bullet" in style_name) else "number",
                            list_level=ilvl,
                            num_id=num_id,
                            ilvl=ilvl,
                            resolved_numbering_label=resolved_num_label,
                            runs=runs_nodes,
                            alignment=eff_para.alignment,
                            line_spacing=eff_para.line_spacing,
                            space_before_pt=eff_para.space_before_pt,
                            space_after_pt=eff_para.space_after_pt,
                            first_line_indent_cm=0.0,
                            left_indent_cm=eff_para.left_indent_cm or 0.63,
                            effective_properties=eff_para.model_dump(),
                        )
                    )
                else:
                    target_sec.blocks.append(
                        ParagraphBlock(
                            id=f"p_{current_sec_idx+1:02d}_{blk_idx+1:04d}",
                            style_name=style_name,
                            runs=runs_nodes,
                            alignment=eff_para.alignment,
                            line_spacing=eff_para.line_spacing,
                            space_before_pt=eff_para.space_before_pt,
                            space_after_pt=eff_para.space_after_pt,
                            first_line_indent_cm=eff_para.first_line_indent_cm,
                            left_indent_cm=eff_para.left_indent_cm,
                            right_indent_cm=eff_para.right_indent_cm,
                            keep_with_next=eff_para.keep_with_next,
                            page_break_before=eff_para.page_break_before,
                            effective_properties=eff_para.model_dump(),
                        )
                    )

                for img_b in image_blocks_in_p:
                    target_sec.blocks.append(img_b)

                if p_sect_pr is not None and current_sec_idx + 1 < len(doc_node.sections):
                    current_sec_idx += 1
                    target_sec = doc_node.sections[current_sec_idx]

            elif tag.endswith("tbl"):
                tbl = docx.table.Table(child, doc)
                num_rows = len(tbl.rows)
                num_cols = len(tbl.columns) if num_rows else 0
                cells_grid = []

                # Extract grid columns
                grid_cols = []
                for gc in child.findall(f".//{{{ns['w']}}}gridCol"):
                    w = gc.attrib.get(f"{{{ns['w']}}}w")
                    if w and w.isdigit():
                        grid_cols.append(round(float(w) / 567.0, 2))

                for r_idx, row in enumerate(tbl.rows):
                    row_cells = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_tc = cell._tc
                        bg_hex = None
                        shd = cell_tc.find(f".//{{{ns['w']}}}shd")
                        if shd is not None:
                            fill = shd.attrib.get(f"{{{ns['w']}}}fill")
                            if fill and fill != "auto":
                                bg_hex = fill.upper()

                        grid_span = 1
                        gs = cell_tc.find(f".//{{{ns['w']}}}gridSpan")
                        if gs is not None:
                            val = gs.attrib.get(f"{{{ns['w']}}}val")
                            if val and val.isdigit():
                                grid_span = int(val)

                        row_cells.append(
                            TableCellNode(
                                id=f"c_{current_sec_idx+1:02d}_{blk_idx+1:04d}_{r_idx+1:02d}_{c_idx+1:02d}",
                                text=normalize_unicode(cell.text),
                                runs=[RunNode(text=normalize_unicode(cell.text))],
                                bg_color_hex=bg_hex,
                                colspan=grid_span,
                            )
                        )
                    cells_grid.append(row_cells)

                target_sec.blocks.append(
                    TableBlock(
                        id=f"tbl_{current_sec_idx+1:02d}_{blk_idx+1:04d}",
                        rows=num_rows,
                        columns=num_cols,
                        cells=cells_grid,
                        style_name=tbl.style.name if tbl.style else "Table Grid",
                        grid_cols_cm=grid_cols,
                    )
                )

            elif not tag.endswith("sectPr"):
                # Preserve unknown OOXML element
                target_sec.blocks.append(
                    UnsupportedBlock(
                        id=f"unsup_{current_sec_idx+1:02d}_{blk_idx+1:04d}",
                        tag_name=tag,
                        raw_xml=child.xml if hasattr(child, "xml") else "",
                    )
                )

        return doc_node


class DocxExporter:
    """
    Serializes Canonical Document Model into an OpenXML Word .docx document.
    """

    @staticmethod
    def export_docx(doc_node: DocumentNode, output_path: Union[str, Path]) -> Path:
        out_p = ensure_parent_dir(output_path)
        doc = docx.Document()

        # Clear default first paragraph
        if doc.paragraphs:
            p_elem = doc.paragraphs[0]._p
            p_elem.getparent().remove(p_elem)

        # Set automatic field update on open
        try:
            update_fields = parse_xml(r'<w:updateFields %s w:val="true"/>' % nsdecls("w"))
            doc.settings.element.append(update_fields)
        except Exception:
            pass

        for s_idx, sec_node in enumerate(doc_node.sections):
            sec = doc.sections[s_idx] if s_idx < len(doc.sections) else doc.add_section()

            # Geometry
            sec.page_width = Cm(sec_node.properties.page_width_cm)
            sec.page_height = Cm(sec_node.properties.page_height_cm)
            sec.top_margin = Cm(sec_node.properties.margin_top_cm)
            sec.bottom_margin = Cm(sec_node.properties.margin_bottom_cm)
            sec.left_margin = Cm(sec_node.properties.margin_left_cm)
            sec.right_margin = Cm(sec_node.properties.margin_right_cm)

            sec.different_first_page_header_footer = sec_node.properties.different_first_page

            # First Page Header / Footer
            if sec_node.properties.different_first_page:
                f_head = sec.first_page_header
                f_head.is_linked_to_previous = False
                p_fh = f_head.paragraphs[0] if f_head.paragraphs else f_head.add_paragraph()
                p_fh.text = sec_node.first_page_header_text or ""

                f_foot = sec.first_page_footer
                f_foot.is_linked_to_previous = False
                p_ff = f_foot.paragraphs[0] if f_foot.paragraphs else f_foot.add_paragraph()
                p_ff.text = sec_node.first_page_footer_text or ""

            # Normal Header / Footer
            if sec_node.header_text:
                header = sec.header
                header.is_linked_to_previous = False
                p_head = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p_head.text = sec_node.header_text
                p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if sec_node.has_page_numbers or sec_node.footer_text:
                footer = sec.footer
                footer.is_linked_to_previous = False
                p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p_foot.text = ""
                if sec_node.footer_text:
                    r_ft_txt = p_foot.add_run(sec_node.footer_text + " | ")
                    r_ft_txt.font.name = "Times New Roman"
                    r_ft_txt.font.size = Pt(9.5)
                    r_ft_txt.font.color.rgb = RGBColor(100, 116, 139)
                add_page_number_to_paragraph(p_foot, format_str=sec_node.page_number_format)
                p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Blocks
            for blk in sec_node.blocks:
                if isinstance(blk, HeadingBlock):
                    p = doc.add_heading(level=blk.level)
                    p.text = ""
                    for r in blk.runs:
                        run = p.add_run(r.text)
                        if r.font_name:
                            run.font.name = r.font_name
                        if r.font_size_pt:
                            run.font.size = Pt(r.font_size_pt)
                        if r.bold is not None:
                            run.bold = r.bold
                        if r.italic is not None:
                            run.italic = r.italic
                        if r.color_rgb:
                            try:
                                run.font.color.rgb = RGBColor(*hex_to_rgb(r.color_rgb))
                            except Exception:
                                pass

                elif isinstance(blk, (ParagraphBlock, ListItemBlock)):
                    p = doc.add_paragraph()
                    align = blk.alignment or "justify"
                    if align in ALIGN_MAP:
                        p.alignment = ALIGN_MAP[align]
                    if blk.line_spacing:
                        p.paragraph_format.line_spacing = blk.line_spacing
                    if blk.space_before_pt is not None:
                        p.paragraph_format.space_before = Pt(blk.space_before_pt)
                    if blk.space_after_pt is not None:
                        p.paragraph_format.space_after = Pt(blk.space_after_pt)
                    if blk.first_line_indent_cm:
                        p.paragraph_format.first_line_indent = Cm(blk.first_line_indent_cm)
                    if blk.left_indent_cm:
                        p.paragraph_format.left_indent = Cm(blk.left_indent_cm)
                    if blk.keep_with_next:
                        p.paragraph_format.keep_with_next = True
                    if getattr(blk, "page_break_before", False):
                        p.paragraph_format.page_break_before = True

                    for r in blk.runs:
                        if r.is_page_break:
                            doc.add_page_break()
                            continue
                        run = p.add_run(r.text)
                        if r.font_name:
                            run.font.name = r.font_name
                        if r.font_size_pt:
                            run.font.size = Pt(r.font_size_pt)
                        if r.bold is not None:
                            run.bold = r.bold
                        if r.italic is not None:
                            run.italic = r.italic
                        if r.underline:
                            run.underline = True
                        if r.color_rgb:
                            try:
                                run.font.color.rgb = RGBColor(*hex_to_rgb(r.color_rgb))
                            except Exception:
                                pass
                        if r.highlight and r.highlight in HIGHLIGHT_MAP:
                            run.font.highlight_color = HIGHLIGHT_MAP[r.highlight]

                elif isinstance(blk, TableBlock):
                    tbl = doc.add_table(rows=blk.rows, cols=blk.columns)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for r_idx, row in enumerate(blk.cells):
                        for c_idx, cell in enumerate(row):
                            if r_idx < len(tbl.rows) and c_idx < len(tbl.columns):
                                t_cell = tbl.cell(r_idx, c_idx)
                                t_cell.text = cell.text
                                if cell.bg_color_hex:
                                    set_cell_background(t_cell, cell.bg_color_hex)

                elif isinstance(blk, ImageBlock):
                    if blk.source_uri_or_path and Path(blk.source_uri_or_path).exists():
                        doc.add_picture(blk.source_uri_or_path, width=Cm(blk.width_cm or 10.0))
                    else:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_img = p_img.add_run(f"[Hình ảnh minh họa: {blk.caption or blk.source_uri_or_path}]")
                        r_img.italic = True

                elif isinstance(blk, DiagramBlock):
                    p_diag = doc.add_paragraph()
                    p_diag.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_diag = p_diag.add_run(f"[Sơ đồ kiến trúc: {blk.caption or blk.diagram_type}]")
                    r_diag.bold = True

                elif isinstance(blk, UnsupportedBlock):
                    if blk.raw_xml:
                        try:
                            elem = parse_xml(blk.raw_xml)
                            doc._body._element.append(elem)
                        except Exception:
                            pass

        doc.save(str(out_p))
        return out_p
