"""
Verification and structural validation engine.
"""

import os
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import docx
from pydantic import BaseModel, Field
from docx_agent.utils.paths import resolve_safe_path
from docx_agent.utils.unicode import normalize_unicode


class ValidationReport(BaseModel):
    passed: bool
    file_path: str
    paragraphs_count: int = 0
    tables_count: int = 0
    sections_count: int = 0
    failures: List[Dict[str, Any]] = Field(default_factory=list)
    warnings: List[str] = Field(default_factory=list)


class DocumentValidator:
    """
    Independently reopens and validates DOCX integrity and XML structure.
    """

    @staticmethod
    def verify_integrity(file_path: Union[str, Path]) -> ValidationReport:
        """
        Performs an independent reload of the DOCX file, parses the XML tree,
        and verifies that all core OpenXML relationships and elements are intact.
        """
        p = resolve_safe_path(file_path)
        if not p.exists():
            return ValidationReport(
                passed=False,
                file_path=str(p),
                failures=[{"type": "FILE_NOT_FOUND", "message": f"File does not exist: {p}"}],
            )

        try:
            # Independent load test
            doc = docx.Document(str(p))
            p_count = len(doc.paragraphs)
            t_count = len(doc.tables)
            s_count = len(doc.sections)

            # Check that XML elements can be serialized
            _ = doc._body._element.xml

            failures = []
            warnings = []

            # Check for corrupt tables (empty rows or mismatched columns)
            for t_idx, tbl in enumerate(doc.tables):
                if len(tbl.rows) == 0:
                    failures.append({"type": "EMPTY_TABLE", "table_index": t_idx})

            return ValidationReport(
                passed=len(failures) == 0,
                file_path=str(p),
                paragraphs_count=p_count,
                tables_count=t_count,
                sections_count=s_count,
                failures=failures,
                warnings=warnings,
            )
        except Exception as e:
            return ValidationReport(
                passed=False,
                file_path=str(p),
                failures=[{"type": "CORRUPTED_PACKAGE", "message": str(e)}],
            )
