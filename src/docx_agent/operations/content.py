"""
High-level content operations for document manipulation.
"""

from typing import Optional, Union, Dict, Any, List
import docx
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.ooxml.runs import replace_in_paragraph_preserving_formatting
from docx_agent.core.exceptions import ElementNotFoundError


class ContentOperations:
    """
    Handles safe, format-preserving content mutations.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def replace_text(
        self,
        target: str,
        replacement: str,
        scope_target: Optional[Union[str, int, Dict[str, Any]]] = None,
        count: Optional[int] = None,
    ) -> int:
        """
        Replaces occurrences of `target` text with `replacement`.
        If `scope_target` is provided, restricts replacement to that element.
        Otherwise replaces across all paragraphs and tables in the document.
        """
        total_replaced = 0

        if scope_target is not None:
            # Single or scoped paragraphs
            paragraphs = self.resolver.resolve_paragraphs(scope_target)
            for p in paragraphs:
                n = replace_in_paragraph_preserving_formatting(p, target, replacement, count)
                total_replaced += n
                if count and total_replaced >= count:
                    break
        else:
            # Entire document paragraphs
            for p in self.model.doc.paragraphs:
                n = replace_in_paragraph_preserving_formatting(p, target, replacement, count)
                total_replaced += n
                if count and total_replaced >= count:
                    break

            # Entire document tables
            if count is None or total_replaced < count:
                for table in self.model.doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                n = replace_in_paragraph_preserving_formatting(p, target, replacement, count)
                                total_replaced += n
                                if count and total_replaced >= count:
                                    break

        self.model.reindex()
        return total_replaced

    def insert_paragraph(
        self,
        text: str,
        target: Union[str, int, Dict[str, Any]],
        position: str = "after",  # "before" or "after"
        style: Optional[str] = None,
    ) -> str:
        """Inserts a new paragraph before or after a target paragraph."""
        paragraphs = self.resolver.resolve_paragraphs(target, single=True)
        ref_p = paragraphs[0]

        # Use python-docx insert_paragraph_before
        if position == "before":
            new_p = ref_p.insert_paragraph_before(text, style=style)
        else:
            # Insert after: insert before ref_p's next sibling or append
            parent = ref_p._p.getparent()
            new_p_elem = OxmlElement("w:p")
            ref_p._p.addnext(new_p_elem)
            new_p = Paragraph(new_p_elem, ref_p._parent)
            if text:
                new_p.add_run(text)
            if style:
                new_p.style = style

        self.model.reindex()
        return self.model.identity.get_paragraph_id(new_p) or "unknown"

    def append_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        heading_level: Optional[int] = None,
    ) -> str:
        """Appends a paragraph or heading at the end of the document."""
        if heading_level is not None:
            p = self.model.doc.add_heading(text, level=heading_level)
        elif style and style.startswith("Heading"):
            try:
                lvl = int("".join(filter(str.isdigit, style)) or 1)
                p = self.model.doc.add_heading(text, level=lvl)
            except ValueError:
                p = self.model.doc.add_paragraph(text, style=style)
        else:
            p = self.model.doc.add_paragraph(text, style=style)

        self.model.reindex()
        return self.model.identity.get_paragraph_id(p) or "unknown"

    def prepend_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
    ) -> str:
        """Prepends a paragraph at the very beginning of the document."""
        if self.model.doc.paragraphs:
            first_p = self.model.doc.paragraphs[0]
            new_p = first_p.insert_paragraph_before(text, style=style)
        else:
            new_p = self.model.doc.add_paragraph(text, style=style)

        self.model.reindex()
        return self.model.identity.get_paragraph_id(new_p) or "unknown"

    def delete_element(self, target: Union[str, int, Dict[str, Any]]) -> bool:
        """Deletes a target paragraph or table from the document."""
        if isinstance(target, str) and target.startswith("tbl_"):
            tbl = self.resolver.resolve_table(target)
            parent = tbl._tbl.getparent()
            parent.remove(tbl._tbl)
            self.model.reindex()
            return True
        else:
            paragraphs = self.resolver.resolve_paragraphs(target, single=True)
            p = paragraphs[0]
            parent = p._p.getparent()
            parent.remove(p._p)
            self.model.reindex()
            return True
