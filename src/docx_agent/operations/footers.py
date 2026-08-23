"""
Footer operations for document sections including dynamic page numbering fields.
"""

from typing import Optional, Union, Dict, Any, List
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.ooxml.fields import add_page_number_to_paragraph, add_field_simple
from docx_agent.operations.paragraphs import ALIGN_MAP


class FooterOperations:
    """
    Manages document footers and dynamic page number fields.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def set_footer(
        self,
        text: str,
        target_section: Optional[Union[str, int]] = None,
        alignment: str = "center",
        unlink_from_previous: bool = True,
    ) -> bool:
        """Sets plain text footer on section."""
        sections = [self.resolver.resolve_section(target_section)] if target_section is not None else list(self.model.doc.sections)

        for sec in sections:
            footer = sec.footer
            if unlink_from_previous:
                footer.is_linked_to_previous = False

            if not footer.paragraphs:
                p = footer.add_paragraph()
            else:
                p = footer.paragraphs[0]
                p.text = ""

            p.text = text
            align_enum = ALIGN_MAP.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            p.alignment = align_enum

        return True

    def set_page_numbers(
        self,
        target_section: Optional[Union[str, int]] = None,
        format_str: str = "Page {PAGE} of {NUMPAGES}",
        alignment: str = "center",
        unlink_from_previous: bool = True,
    ) -> bool:
        """
        Inserts dynamic Word page number fields into the footer.
        """
        sections = [self.resolver.resolve_section(target_section)] if target_section is not None else list(self.model.doc.sections)

        for sec in sections:
            footer = sec.footer
            if unlink_from_previous:
                footer.is_linked_to_previous = False

            if not footer.paragraphs:
                p = footer.add_paragraph()
            else:
                p = footer.paragraphs[0]
                p.text = ""

            add_page_number_to_paragraph(p, format_str=format_str)
            align_enum = ALIGN_MAP.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            p.alignment = align_enum

        return True

    def clear_footer(self, target_section: Optional[Union[str, int]] = None) -> bool:
        """Clears footer for section(s)."""
        sections = [self.resolver.resolve_section(target_section)] if target_section is not None else list(self.model.doc.sections)
        for sec in sections:
            footer = sec.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.text = ""
        return True
