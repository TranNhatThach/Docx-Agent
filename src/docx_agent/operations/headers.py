"""
Header operations for document sections.
"""

from typing import Optional, Union, Dict, Any, List
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.ooxml.fields import add_field_simple
from docx_agent.operations.paragraphs import ALIGN_MAP


class HeaderOperations:
    """
    Manages document headers across sections.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def set_header(
        self,
        text: str,
        target_section: Optional[Union[str, int]] = None,
        alignment: str = "right",
        unlink_from_previous: bool = True,
    ) -> bool:
        """Sets header content on target section (or first section if None)."""
        sections = [self.resolver.resolve_section(target_section)] if target_section is not None else list(self.model.doc.sections)

        for sec in sections:
            header = sec.header
            if unlink_from_previous:
                header.is_linked_to_previous = False

            # Clear existing paragraphs in header
            if not header.paragraphs:
                p = header.add_paragraph()
            else:
                p = header.paragraphs[0]
                p.text = ""

            p.text = text
            align_enum = ALIGN_MAP.get(alignment.lower(), WD_ALIGN_PARAGRAPH.RIGHT)
            p.alignment = align_enum

        return True

    def clear_header(self, target_section: Optional[Union[str, int]] = None) -> bool:
        """Clears header content for section(s)."""
        sections = [self.resolver.resolve_section(target_section)] if target_section is not None else list(self.model.doc.sections)
        for sec in sections:
            header = sec.header
            header.is_linked_to_previous = False
            for p in header.paragraphs:
                p.text = ""
        return True
