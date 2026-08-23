"""
Image operations: Insertion, sizing, alignment, captions, and relationship inspection.
"""

from pathlib import Path
from typing import Optional, Union, Dict, Any, List
import docx
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.core.exceptions import ElementNotFoundError, DocxAgentError, ErrorCode
from docx_agent.utils.paths import resolve_safe_path
from docx_agent.operations.paragraphs import ALIGN_MAP


class ImageOperations:
    """
    Handles image insertion, scaling, positioning, and captions.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def insert_image(
        self,
        image_path: Union[str, Path],
        target: Optional[Union[str, int, Dict[str, Any]]] = None,
        position: str = "after",  # "after", "before", "append"
        width_cm: Optional[float] = None,
        height_cm: Optional[float] = None,
        alignment: str = "center",
        caption: Optional[str] = None,
    ) -> str:
        """
        Inserts an image into the document with optional dimensions, alignment, and caption.
        """
        img_p = resolve_safe_path(image_path)
        if not img_p.exists():
            raise DocxAgentError(
                message=f"Image file not found: {img_p}",
                code=ErrorCode.IMAGE_ERROR,
                details={"image_path": str(img_p)},
            )

        kwargs = {}
        if width_cm is not None:
            kwargs["width"] = Cm(width_cm)
        if height_cm is not None:
            kwargs["height"] = Cm(height_cm)

        # Target paragraph placement
        if target is not None and position in ("before", "after"):
            ref_paragraphs = self.resolver.resolve_paragraphs(target, single=True)
            ref_p = ref_paragraphs[0]
            if position == "before":
                p = ref_p.insert_paragraph_before()
            else:
                p = self.model.doc.add_paragraph()  # default fallback or insert after
                ref_p._p.addnext(p._p)
        else:
            p = self.model.doc.add_paragraph()

        # Alignment
        align_enum = ALIGN_MAP.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
        p.alignment = align_enum

        # Add image run
        run = p.add_run()
        run.add_picture(str(img_p), **kwargs)

        # Add Caption if requested
        if caption:
            cap_p = p.insert_paragraph_before() if position == "before" else self.model.doc.add_paragraph()
            if position == "after":
                p._p.addnext(cap_p._p)
            cap_p.text = caption
            try:
                cap_p.style = "Caption"
            except Exception:
                pass
            cap_p.alignment = align_enum

        self.model.reindex()
        return self.model.identity.get_paragraph_id(p) or "unknown"
