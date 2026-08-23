"""
Table operations for creating, modifying, styling, and formatting Word tables.
"""

from typing import Optional, Union, Dict, Any, List
import docx
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.table import Table, _Cell
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.core.exceptions import ElementNotFoundError
from docx_agent.ooxml.tables import (
    set_cell_background,
    set_row_repeat_header,
    set_row_cant_split,
    set_table_borders,
)
from docx_agent.operations.formatting import hex_to_rgb


TABLE_ALIGN_MAP = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}


class TableOperations:
    """
    Manages structured table creation, manipulation, and styling.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def create_table(
        self,
        rows: int,
        columns: int,
        data: Optional[List[List[str]]] = None,
        style: str = "Table Grid",
        alignment: str = "center",
        repeat_header: bool = True,
        col_widths_cm: Optional[List[float]] = None,
    ) -> str:
        """Creates a new table at the end of the document."""
        tbl = self.model.doc.add_table(rows=rows, cols=columns)
        try:
            tbl.style = style
        except Exception:
            tbl.style = "Table Grid"

        align_enum = TABLE_ALIGN_MAP.get(alignment.lower())
        if align_enum:
            tbl.alignment = align_enum

        if data:
            for r_idx, row_values in enumerate(data):
                if r_idx < len(tbl.rows):
                    row = tbl.rows[r_idx]
                    for c_idx, val in enumerate(row_values):
                        if c_idx < len(row.cells):
                            row.cells[c_idx].text = str(val)

        if repeat_header and tbl.rows:
            set_row_repeat_header(tbl.rows[0], True)
            set_cell_background(tbl.rows[0].cells[0], "F2F2F2") if tbl.rows[0].cells else None

        for r in tbl.rows:
            set_row_cant_split(r, True)

        if col_widths_cm:
            for col_idx, width in enumerate(col_widths_cm):
                for row in tbl.rows:
                    if col_idx < len(row.cells):
                        row.cells[col_idx].width = Cm(width)

        self.model.reindex()
        return self.model.identity.get_table_id(tbl) or "unknown"

    def edit_cell(
        self,
        cell_id: str,
        text: str,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color_rgb: Optional[str] = None,
        bg_color_hex: Optional[str] = None,
    ) -> bool:
        """Edits content and formatting of a specific cell."""
        cell = self.resolver.resolve_cell(cell_id)
        cell.text = text
        if (bold is not None or italic is not None or color_rgb is not None) and cell.paragraphs:
            for p in cell.paragraphs:
                for r in p.runs:
                    if bold is not None:
                        r.bold = bold
                    if italic is not None:
                        r.italic = italic
                    if color_rgb is not None:
                        r.font.color.rgb = hex_to_rgb(color_rgb)

        if bg_color_hex:
            set_cell_background(cell, bg_color_hex)

        return True

    def add_row(self, target: Union[str, int], row_data: Optional[List[str]] = None) -> int:
        """Appends a new row to the table."""
        tbl = self.resolver.resolve_table(target)
        row = tbl.add_row()
        set_row_cant_split(row, True)
        if row_data:
            for c_idx, val in enumerate(row_data):
                if c_idx < len(row.cells):
                    row.cells[c_idx].text = str(val)
        self.model.reindex()
        return len(tbl.rows)

    def delete_row(self, target: Union[str, int], row_index: int) -> bool:
        """Deletes a row from table by row index."""
        tbl = self.resolver.resolve_table(target)
        if 0 <= row_index < len(tbl.rows):
            row_elem = tbl.rows[row_index]._tr
            row_elem.getparent().remove(row_elem)
            self.model.reindex()
            return True
        return False

    def set_borders(self, target: Union[str, int], color: str = "D3D3D3", sz: str = "4") -> bool:
        """Sets custom borders on table."""
        tbl = self.resolver.resolve_table(target)
        set_table_borders(tbl, color=color, sz=sz)
        return True
