"""
Field and Table of Contents (TOC) operations.
"""

from typing import Optional, Union, Dict, Any, List
import docx
from docx.text.paragraph import Paragraph
from docx_agent.core.document import DocumentModel
from docx_agent.core.resolver import TargetResolver
from docx_agent.ooxml.fields import add_toc_field, add_field_simple


class FieldOperations:
    """
    Manages insertion of TOC and general OOXML fields.
    """

    def __init__(self, model: DocumentModel):
        self.model = model
        self.resolver = TargetResolver(model)

    def insert_toc(
        self,
        target: Optional[Union[str, int, Dict[str, Any]]] = None,
        position: str = "before",  # "before", "after", "prepend"
        levels: str = "1-3",
        title: str = "Mục Lục",
    ) -> str:
        """
        Inserts a dynamic Word Table of Contents field into the document.
        """
        if target is not None:
            paragraphs = self.resolver.resolve_paragraphs(target, single=True)
            ref_p = paragraphs[0]
            if position == "before":
                title_p = ref_p.insert_paragraph_before(title, style="Heading 1")
                toc_p = ref_p.insert_paragraph_before()
            else:
                title_p = self.model.doc.add_paragraph(title, style="Heading 1")
                ref_p._p.addnext(title_p._p)
                toc_p = self.model.doc.add_paragraph()
                title_p._p.addnext(toc_p._p)
        else:
            # Prepend or append at top
            if self.model.doc.paragraphs:
                first_p = self.model.doc.paragraphs[0]
                title_p = first_p.insert_paragraph_before(title, style="Heading 1")
                toc_p = first_p.insert_paragraph_before()
            else:
                title_p = self.model.doc.add_paragraph(title, style="Heading 1")
                toc_p = self.model.doc.add_paragraph()

        add_toc_field(toc_p, levels=levels)
        self.model.reindex()
        return self.model.identity.get_paragraph_id(toc_p) or "unknown"
