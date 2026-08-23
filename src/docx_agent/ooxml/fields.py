"""
OOXML field manipulation utilities for Word dynamic fields:
PAGE, NUMPAGES, DATE, TOC, TITLE, AUTHOR.
"""

from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def add_field_simple(paragraph: Paragraph, instr_text: str, default_display: str = "1") -> None:
    """
    Appends a dynamic w:fldSimple field to a paragraph (e.g. PAGE, NUMPAGES, DATE).
    """
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="{instr_text}"><w:r><w:rPr/><w:t>{default_display}</w:t></w:r></w:fldSimple>'
    fld_elem = parse_xml(fld_xml)
    paragraph._p.append(fld_elem)


def add_page_number_to_paragraph(paragraph: Paragraph, format_str: str = "Page {PAGE} of {NUMPAGES}") -> None:
    """
    Adds a formatted page numbering field structure to a paragraph.
    Handles 'Page {PAGE} of {NUMPAGES}' or just '{PAGE}'.
    """
    parts = format_str.split("{PAGE}")
    if len(parts) == 1:
        # Just plain text or unsupported template
        paragraph.add_run(format_str)
        return

    # Prefix
    if parts[0]:
        paragraph.add_run(parts[0])

    # PAGE field
    add_field_simple(paragraph, "PAGE", "1")

    # Suffix or intermediate
    after_page = parts[1]
    if "{NUMPAGES}" in after_page:
        sub_parts = after_page.split("{NUMPAGES}")
        if sub_parts[0]:
            paragraph.add_run(sub_parts[0])
        add_field_simple(paragraph, "NUMPAGES", "1")
        if len(sub_parts) > 1 and sub_parts[1]:
            paragraph.add_run(sub_parts[1])
    else:
        if after_page:
            paragraph.add_run(after_page)


def add_toc_field(paragraph: Paragraph, levels: str = "1-3") -> None:
    """
    Adds a dynamic Table of Contents field to a paragraph.
    """
    # Complex TOC field structure for maximum Word compatibility
    toc_xml = f"""
    <w:p {nsdecls("w")}>
        <w:pPr>
            <w:pStyle w:val="TOCHeading"/>
        </w:pPr>
        <w:fldSimple w:instr="TOC \\o &quot;{levels}&quot; \\h \\z \\u">
            <w:r>
                <w:rPr/>
                <w:t>Table of Contents (Press F9 or right-click to update in Word)</w:t>
            </w:r>
        </w:fldSimple>
    </w:p>
    """
    elem = parse_xml(toc_xml)
    paragraph._p.getparent().replace(paragraph._p, elem)
